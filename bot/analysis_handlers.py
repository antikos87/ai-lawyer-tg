#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Обработчики для анализа юридических документов
"""

import logging
import os
import tempfile
import re
import asyncio
from enum import Enum
from typing import Dict, Any, List, Optional, Tuple
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, Document
from telegram.ext import ContextTypes, ConversationHandler
from telegram.constants import ChatAction

# Импорты для извлечения текста
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pytesseract
from PIL import Image
import PyPDF2

from ai_gigachat.client import gigachat_client
from bot.middleware import subscription_required, add_usage_info_to_response

logger = logging.getLogger(__name__)


class AnalysisStates(Enum):
    """Состояния диалога анализа документов"""
    DOCUMENT_UPLOAD = 0
    ANALYSIS_TYPE_SELECTION = 1
    TEXT_PROCESSING = 2
    ANALYSIS_PROCESSING = 3
    RESULTS_REVIEW = 4
    ADDITIONAL_ACTIONS = 5
    MULTIPLE_IMAGES = 6  # Новое состояние для загрузки нескольких изображений


# Поддерживаемые типы файлов
SUPPORTED_EXTENSIONS = {
    'document': ['.doc', '.docx', '.pdf'],
    'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']
}

# Максимальный размер файла (10 МБ)
MAX_FILE_SIZE = 10 * 1024 * 1024

# Типы анализа (соответствуют ключам в GigaChat клиенте)
ANALYSIS_TYPES = {
    'document_summary': {
        'name': '📋 Краткое описание',
        'description': 'Краткое описание содержания и назначения документа',
        'icon': '📋'
    },
    'law_compliance': {
        'name': '⚖️ Проверка соответствия закону',
        'description': 'Анализ соответствия документа действующему законодательству РФ',
        'icon': '⚖️'
    },
    'error_detection': {
        'name': '🔍 Найти ошибки',
        'description': 'Поиск юридических, технических и логических ошибок в документе',
        'icon': '🔍'
    },
    'risk_assessment': {
        'name': '⚠️ Оценить риски',
        'description': 'Анализ правовых рисков и потенциальных проблем',
        'icon': '⚠️'
    },
    'recommendations': {
        'name': '💡 Дать рекомендации',
        'description': 'Предложения по улучшению и оптимизации документа',
        'icon': '💡'
    },
    'correspondence_analysis': {
        'name': '📧 Анализировать переписку',
        'description': 'Специальный анализ деловой переписки и коммуникаций',
        'icon': '📧'
    }
}


async def analyze_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработчик команды /analyze - начало анализа документов
    """
    user = update.effective_user
    user_name = user.first_name if user.first_name else "пользователь"
    
    # Очищаем предыдущие данные
    context.user_data.clear()
    
    welcome_text = (
        f"📊 **Анализ юридических документов**\n\n"
        f"Привет, {user_name}! Я помогу проанализировать ваш документ.\n\n"
        
        "📋 **Поддерживаемые форматы:**\n"
        "• **Документы:** DOC, DOCX, PDF\n"
        "• **Изображения:** JPG, PNG (сканы документов)\n"
        "• **Максимальный размер:** 10 МБ\n\n"
        
        "🎯 **Типы анализа:**\n"
        "• Проверка соответствия закону\n"
        "• Поиск ошибок и недочетов\n"
        "• Оценка правовых рисков\n"
        "• Рекомендации по улучшению\n"
        "• Анализ деловой переписки\n\n"
        
        "📎 **Загрузите документ для анализа:**"
    )
    
    keyboard = [
        [InlineKeyboardButton("❌ Отмена", callback_data="cancel_analysis")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    # Определяем откуда пришел запрос
    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.reply_text(
            welcome_text,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
    else:
        await update.message.reply_text(
            welcome_text,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
    
    return AnalysisStates.DOCUMENT_UPLOAD.value


@subscription_required('analysis', 'анализ документа')
async def handle_document_upload(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработчик загрузки документа с поддержкой множественных изображений
    """
    logger.info(f"Начало обработки загрузки документа от пользователя {update.effective_user.id}")
    
    # Показываем индикатор обработки
    await update.message.chat.send_action(ChatAction.TYPING)
    
    document = update.message.document
    photo = update.message.photo
    
    # Проверяем что загружен файл
    if not document and not photo:
        await update.message.reply_text(
            "❌ **Ошибка загрузки**\n\n"
            "Пожалуйста, загрузите документ в поддерживаемом формате:\n"
            "• DOC, DOCX, PDF - документы\n"
            "• JPG, PNG - изображения (сканы)\n\n"
            "Максимальный размер файла: 10 МБ",
            parse_mode='Markdown'
        )
        return AnalysisStates.DOCUMENT_UPLOAD.value
    
    # Обработка обычного документа
    if document:
        file_name = document.file_name
        file_size = document.file_size
        file_id = document.file_id
        
        # Проверяем размер файла
        if file_size > MAX_FILE_SIZE:
            await update.message.reply_text(
                f"❌ **Файл слишком большой**\n\n"
                f"Размер файла: {file_size / (1024*1024):.1f} МБ\n"
                f"Максимальный размер: 10 МБ\n\n"
                "Пожалуйста, загрузите файл меньшего размера.",
                parse_mode='Markdown'
            )
            return AnalysisStates.DOCUMENT_UPLOAD.value
        
        # Проверяем расширение файла
        file_extension = os.path.splitext(file_name.lower())[1]
        
        # Если это изображение, обрабатываем как изображение
        if file_extension in SUPPORTED_EXTENSIONS['image']:
            file_type = 'image'
        elif file_extension in SUPPORTED_EXTENSIONS['document']:
            file_type = 'document'
        else:
            await update.message.reply_text(
                f"❌ **Неподдерживаемый формат**\n\n"
                f"Файл: `{file_name}`\n"
                f"Формат: `{file_extension}`\n\n"
                "Поддерживаемые форматы:\n"
                "• **Документы:** DOC, DOCX, PDF\n"
                "• **Изображения:** JPG, PNG, BMP, TIFF\n\n"
                "Пожалуйста, загрузите файл в поддерживаемом формате.",
                parse_mode='Markdown'
            )
            return AnalysisStates.DOCUMENT_UPLOAD.value
        
    # Обработка изображения
    else:  # photo
        logger.info(f"Обработка изображения от пользователя {update.effective_user.id}")
        
        # Берем изображение наилучшего качества
        photo_file = photo[-1]
        file_size = photo_file.file_size
        file_id = photo_file.file_id
        file_name = f"image_{photo_file.file_id}.jpg"
        file_extension = '.jpg'
        file_type = 'image'
        
        # Проверяем размер
        if file_size > MAX_FILE_SIZE:
            logger.warning(f"Изображение слишком большое: {file_size} байт от пользователя {update.effective_user.id}")
            await update.message.reply_text(
                f"❌ **Изображение слишком большое**\n\n"
                f"Размер: {file_size / (1024*1024):.1f} МБ\n"
                f"Максимальный размер: 10 МБ\n\n"
                "Пожалуйста, загрузите изображение меньшего размера.",
                parse_mode='Markdown'
            )
            return AnalysisStates.DOCUMENT_UPLOAD.value
        
        # Проверяем, есть ли уже загруженные изображения для объединения
        if 'uploaded_images' not in context.user_data:
            context.user_data['uploaded_images'] = []
        
        # Добавляем текущее изображение к списку
        context.user_data['uploaded_images'].append({
            'file_id': file_id,
            'file_name': file_name,
            'file_size': file_size,
            'file_extension': file_extension,
            'file_type': file_type
        })
        
        # Если это первое изображение, предлагаем добавить еще
        if len(context.user_data['uploaded_images']) == 1:
            keyboard = [
                [InlineKeyboardButton("📷 Добавить еще изображение", callback_data="add_more_images")],
                [InlineKeyboardButton("✅ Продолжить с одним изображением", callback_data="process_single_image")],
                [InlineKeyboardButton("❌ Отмена", callback_data="cancel_analysis")]
            ]
            
            await update.message.reply_text(
                f"📷 **Изображение загружено!**\n\n"
                f"📊 **Размер:** {file_size / (1024*1024):.1f} МБ\n\n"
                "💡 **Совет:** Если договор состоит из нескольких страниц, "
                "вы можете загрузить все изображения подряд для лучшего распознавания текста.\n\n"
                "**Что делаем дальше?**",
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode='Markdown'
            )
            return AnalysisStates.MULTIPLE_IMAGES.value
        
        # Если уже есть изображения, добавляем к коллекции
        else:
            total_images = len(context.user_data['uploaded_images'])
            total_size = sum(img['file_size'] for img in context.user_data['uploaded_images'])
            
            keyboard = [
                [InlineKeyboardButton("📷 Добавить еще изображение", callback_data="add_more_images")],
                [InlineKeyboardButton("✅ Обработать все изображения", callback_data="process_all_images")],
                [InlineKeyboardButton("❌ Отмена", callback_data="cancel_analysis")]
            ]
            
            await update.message.reply_text(
                f"📷 **Изображение {total_images} добавлено!**\n\n"
                f"📊 **Всего изображений:** {total_images}\n"
                f"📊 **Общий размер:** {total_size / (1024*1024):.1f} МБ\n\n"
                "**Продолжаем добавлять или обрабатываем?**",
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode='Markdown'
            )
            return AnalysisStates.MULTIPLE_IMAGES.value
    
    # Сохраняем информацию о файле
    context.user_data['file_info'] = {
        'file_id': file_id,
        'file_name': file_name,
        'file_size': file_size,
        'file_extension': file_extension,
        'file_type': file_type
    }
    
    # Показываем индикатор обработки
    processing_message = await update.message.reply_text(
        "🔄 **Обрабатываю документ...**\n\n"
        "📥 Загружаю файл...",
        parse_mode='Markdown'
    )
    
    try:
        # Обновляем статус
        await processing_message.edit_text(
            "🔄 **Обрабатываю документ...**\n\n"
            f"📄 {'🖼️ Сканирую изображение...' if file_type == 'image' else '📝 Извлекаю текст из файла...'}\n"
            "⏳ *Это может занять несколько секунд*",
            parse_mode='Markdown'
        )
        
        # Извлекаем текст из файла
        extracted_text = await extract_text_from_file(context.user_data['file_info'])
        
        # Для изображений (OCR) проверяем менее строго
        min_length = 3 if file_type == 'image' else 10
        
        if not extracted_text or len(extracted_text.strip()) < min_length:
            # Если это изображение и OCR вернул предупреждение, все равно продолжаем
            if file_type == 'image' and extracted_text and "OCR распознал очень мало текста" in extracted_text:
                pass  # Продолжаем с предупреждением
            else:
                # Создаем клавиатуру с навигацией для ошибки
                error_keyboard = [
                    [InlineKeyboardButton("🔄 Попробовать еще раз", callback_data="retry_upload")],
                    [InlineKeyboardButton("🔙 Назад к выбору файла", callback_data="back_to_upload")],
                    [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
                ]
                
                await processing_message.edit_text(
                    "❌ **Ошибка обработки файла**\n\n"
                    "Не удалось извлечь текст из документа.\n\n"
                    "**Возможные причины:**\n"
                    "• Плохое качество скана\n"
                    "• Поврежденный файл\n"
                    "• Документ не содержит текста\n\n"
                    "**Что можно сделать:**\n"
                    "• Попробовать загрузить файл заново\n"
                    "• Использовать изображение лучшего качества\n"
                    "• Попробовать другой формат файла",
                    reply_markup=InlineKeyboardMarkup(error_keyboard),
                    parse_mode='Markdown'
                )
                return AnalysisStates.DOCUMENT_UPLOAD.value
        
        # Сохраняем извлеченный текст
        context.user_data['document_text'] = extracted_text
        
        # Финальное обновление с результатом
        file_icon = "🖼️" if file_type == 'image' else "📄"
        success_text = "✅ **Документ обработан успешно!**\n\n"
        
        # Безопасное имя файла для Markdown
        safe_file_name = escape_markdown(file_name)
        
        if file_type == 'image':
            success_text += f"{file_icon} **Изображение:** `{safe_file_name}`\n"
            success_text += f"📊 **Размер:** {file_size / 1024 / 1024:.1f} МБ\n"
            success_text += f"🔤 **Текст распознан:** {len(extracted_text)} символов\n"
            if "OCR распознал очень мало текста" in extracted_text:
                success_text += "\n⚠️ *Распознано мало текста, но анализ возможен*"
        else:
            success_text += f"{file_icon} **Файл:** `{safe_file_name}`\n"
            success_text += f"📊 **Размер:** {file_size / 1024 / 1024:.1f} МБ\n"
            success_text += f"📝 **Текст извлечен:** {len(extracted_text)} символов"
        
        try:
            await processing_message.edit_text(
                success_text,
                parse_mode='Markdown'
            )
        except Exception as edit_error:
            logger.warning(f"Ошибка обновления сообщения с Markdown: {edit_error}")
            # Пытаемся без форматирования
            try:
                await processing_message.edit_text(
                    success_text.replace('**', '').replace('`', '').replace('*', '')
                )
            except Exception as edit_error2:
                logger.error(f"Критическая ошибка обновления сообщения: {edit_error2}")
                # Продолжаем работу даже если не удалось обновить сообщение
        
    except Exception as e:
        # Подробное логирование ошибки
        logger.error(f"Критическая ошибка обработки файла для пользователя {update.effective_user.id}")
        logger.error(f"Тип ошибки: {type(e).__name__}")
        logger.error(f"Сообщение ошибки: {str(e)}")
        logger.error(f"Информация о файле: {context.user_data.get('file_info', 'Нет данных')}")
        
        # Создаем клавиатуру с навигацией для критической ошибки
        error_keyboard = [
            [InlineKeyboardButton("🔄 Попробовать еще раз", callback_data="retry_upload")],
            [InlineKeyboardButton("🔙 Назад к выбору файла", callback_data="back_to_upload")],
            [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
        ]
        
        # Определяем тип ошибки для более точного сообщения
        if "timeout" in str(e).lower() or "time" in str(e).lower():
            error_message = (
                "❌ **Превышено время обработки**\n\n"
                "Файл слишком большой или сложный для обработки.\n\n"
                "**Попробуйте:**\n"
                "• Уменьшить размер файла\n"
                "• Разбить документ на части\n"
                "• Использовать другой формат"
            )
        elif "memory" in str(e).lower() or "размер" in str(e).lower():
            error_message = (
                "❌ **Файл слишком большой**\n\n"
                "Не хватает ресурсов для обработки файла.\n\n"
                "**Попробуйте:**\n"
                "• Уменьшить размер файла\n"
                "• Использовать сжатие\n"
                "• Разбить на несколько частей"
            )
        else:
            error_message = (
                "❌ **Критическая ошибка обработки**\n\n"
                "Произошла непредвиденная ошибка при обработке файла.\n\n"
                "**Попробуйте:**\n"
                "• Загрузить файл заново\n"
                "• Использовать другой формат\n"
                "• Обратиться в поддержку если проблема повторяется"
            )
        
        await processing_message.edit_text(
            error_message,
            reply_markup=InlineKeyboardMarkup(error_keyboard),
            parse_mode='Markdown'
        )
        return AnalysisStates.DOCUMENT_UPLOAD.value
    
    # Предлагаем выбрать тип анализа
    keyboard = create_analysis_keyboard()
    
    # Безопасное имя файла
    safe_file_name = escape_markdown(file_name)
    
    try:
        await update.message.reply_text(
            f"📄 **Документ получен:** `{safe_file_name}`\n\n"
            f"📊 **Размер:** {file_size / 1024 / 1024:.1f} МБ\n"
            f"📝 **Текст извлечен:** {len(context.user_data['document_text'])} символов\n\n"
            "**Выберите тип анализа:**",
            reply_markup=keyboard,
            parse_mode='Markdown'
        )
    except Exception as send_error:
        logger.warning(f"Ошибка отправки сообщения с Markdown: {send_error}")
        # Отправляем без форматирования
        await update.message.reply_text(
            f"📄 Документ получен: {file_name}\n\n"
            f"📊 Размер: {file_size / 1024 / 1024:.1f} МБ\n"
            f"📝 Текст извлечен: {len(context.user_data['document_text'])} символов\n\n"
            "Выберите тип анализа:",
            reply_markup=keyboard
        )
    
    return AnalysisStates.ANALYSIS_TYPE_SELECTION.value


async def handle_analysis_type_selection(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработчик выбора типа анализа
    """
    query = update.callback_query
    await query.answer()
    
    if query.data.startswith("analysis_type_"):
        # Выбор типа анализа
        analysis_type = query.data.replace("analysis_type_", "")
        
        if analysis_type not in ANALYSIS_TYPES:
            await query.answer("❌ Неизвестный тип анализа")
            return
        
        # Сохраняем выбранный тип анализа
        context.user_data['analysis_type'] = analysis_type
        
        # Если это краткое описание, сразу запускаем анализ
        if analysis_type == 'document_summary':
            # НЕ УДАЛЯЕМ предыдущее сообщение! Добавляем новое
            progress_msg = await query.message.reply_text("⏳ Создаю краткое описание документа...")
            
            # Получаем текст документа
            document_text = context.user_data.get('document_text', '')
            if not document_text:
                await query.message.reply_text("❌ Текст документа не найден. Попробуйте загрузить файл заново.")
                return
            
            try:
                # Выполняем анализ
                analysis_result = await gigachat_client.analyze_document(
                    document_text, 
                    analysis_type
                )
                
                # Обновляем прогресс сообщение
                await progress_msg.edit_text("✅ Краткое описание готово!")
                
                # Устанавливаем флаг что краткое описание выполнено
                context.user_data['summary_done'] = True
                
                # Создаем клавиатуру с остальными типами анализа
                keyboard = create_other_analysis_keyboard()
                
                await query.message.reply_text(
                    analysis_result,
                    reply_markup=keyboard,
                    parse_mode='Markdown'
                )
                return
            except Exception as e:
                await query.message.reply_text(
                    "❌ **Ошибка анализа**\n\n"
                    "Произошла ошибка при создании краткого описания.\n"
                    "Попробуйте повторить попытку.",
                    parse_mode='Markdown'
                )
                return
        
        # Для остальных типов анализа
        analysis_info = ANALYSIS_TYPES[analysis_type]
        
        # Если краткое описание уже было сделано, сразу запускаем анализ
        if context.user_data.get('summary_done'):
            # НЕ УДАЛЯЕМ предыдущие сообщения! Добавляем новое
            progress_msg = await query.message.reply_text("⏳ Выполняю анализ документа...")
            
            # Получаем текст документа
            document_text = context.user_data.get('document_text') or context.user_data.get('extracted_text', '')
            if not document_text:
                await query.message.reply_text("❌ Текст документа не найден. Попробуйте загрузить файл заново.")
                return
            
            try:
                # Выполняем анализ
                analysis_result = await gigachat_client.analyze_document(
                    document_text, 
                    analysis_type
                )
                
                logger.info(f"Получен результат анализа, длина: {len(analysis_result)} символов")
                
                # Обновляем прогресс сообщение
                await progress_msg.edit_text("✅ Анализ завершен!")
                
                # Создаем клавиатуру для завершения или выбора другого анализа
                keyboard = [
                    [InlineKeyboardButton("📊 Другой тип анализа", callback_data="back_to_analysis_types")],
                    [InlineKeyboardButton("✅ Завершить", callback_data="finish_analysis")]
                ]
                
                # Отправляем результат с разбивкой на части если нужно
                try:
                    await send_long_message(
                        query.message.chat,
                        analysis_result,
                        reply_markup=InlineKeyboardMarkup(keyboard),
                        parse_mode='Markdown'
                    )
                    logger.info("Результат анализа успешно отправлен пользователю")
                except Exception as send_error:
                    logger.error(f"Ошибка отправки с разбивкой: {send_error}")
                    # Fallback: пытаемся отправить без форматирования
                    try:
                        await send_long_message(
                            query.message.chat,
                            analysis_result,
                            reply_markup=InlineKeyboardMarkup(keyboard)
                        )
                        logger.info("Результат анализа отправлен без Markdown")
                    except Exception as send_error2:
                        logger.error(f"Критическая ошибка отправки: {send_error2}")
                        raise send_error2
                return AnalysisStates.RESULTS_REVIEW.value
            except Exception as e:
                logger.error(f"ОШИБКА при анализе документа: {type(e).__name__}: {str(e)}")
                await query.message.reply_text(
                    "❌ **Ошибка анализа**\n\n"
                    "Произошла ошибка при анализе документа.\n"
                    "Попробуйте повторить попытку или выбрать другой тип анализа.",
                    parse_mode='Markdown'
                )
                return
        
        # Для первичного выбора (без краткого описания) показываем подтверждение
        keyboard = [
            [InlineKeyboardButton("✅ Начать анализ", callback_data="start_analysis")],
            [InlineKeyboardButton("🔙 Выбрать другой тип", callback_data="back_to_analysis_types")],
            [InlineKeyboardButton("❌ Отмена", callback_data="cancel_analysis")]
        ]
        
        # НЕ УДАЛЯЕМ предыдущие сообщения! Добавляем новое
        await query.message.reply_text(
            f"**{analysis_info['icon']} {analysis_info['name']}**\n\n"
            f"{analysis_info['description']}\n\n"
            "Начать анализ?",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )

    elif query.data == "start_analysis":
        # Начать анализ после подтверждения
        progress_msg = await query.message.reply_text("⏳ Выполняю анализ документа...")
        
        # Получаем данные из контекста
        analysis_type = context.user_data.get('analysis_type')
        # Пробуем получить текст из разных переменных (для совместимости)
        document_text = context.user_data.get('document_text') or context.user_data.get('extracted_text', '')
        
        if not analysis_type or not document_text:
            await query.message.reply_text("❌ Данные для анализа не найдены. Попробуйте загрузить файл заново.")
            return
        
        try:
            # Выполняем анализ
            analysis_result = await gigachat_client.analyze_document(
                document_text, 
                analysis_type
            )
            
            logger.info(f"Получен результат анализа start_analysis, длина: {len(analysis_result)} символов")
            
            # Обновляем прогресс сообщение
            await progress_msg.edit_text("✅ Анализ завершен!")
            
            # Создаем клавиатуру для завершения или выбора другого анализа
            keyboard = [
                [InlineKeyboardButton("📊 Другой тип анализа", callback_data="back_to_analysis_types")],
                [InlineKeyboardButton("✅ Завершить", callback_data="finish_analysis")]
            ]
            
            # Отправляем результат с разбивкой на части если нужно
            try:
                await send_long_message(
                    query.message.chat,
                    analysis_result,
                    reply_markup=InlineKeyboardMarkup(keyboard),
                    parse_mode='Markdown'
                )
                logger.info("Результат анализа успешно отправлен")
            except Exception as send_error:
                logger.error(f"Ошибка отправки start_analysis с разбивкой: {send_error}")
                # Fallback: пытаемся отправить без форматирования
                try:
                    await send_long_message(
                        query.message.chat,
                        analysis_result,
                        reply_markup=InlineKeyboardMarkup(keyboard)
                    )
                    logger.info("Результат анализа start_analysis отправлен без Markdown")
                except Exception as send_error2:
                    logger.error(f"Критическая ошибка отправки start_analysis: {send_error2}")
                    raise send_error2
            return AnalysisStates.RESULTS_REVIEW.value
        except Exception as e:
            logger.error(f"ОШИБКА при анализе документа start_analysis: {type(e).__name__}: {str(e)}")
            await query.message.reply_text(
                "❌ **Ошибка анализа**\n\n"
                "Произошла ошибка при анализе документа.\n"
                "Попробуйте:\n"
                "• Повторить анализ через минуту\n"
                "• Выбрать другой тип анализа\n"
                "• Обратиться в поддержку\n\n"
                "Что делаем дальше?",
                parse_mode='Markdown'
            )
            return AnalysisStates.RESULTS_REVIEW.value
    
    elif query.data == "back_to_analysis_types":
        # Вернуться к выбору типа анализа
        # Проверяем, было ли уже выполнено краткое описание
        if context.user_data.get('summary_done'):
            # Если краткое описание уже было, показываем остальные типы
            keyboard = create_other_analysis_keyboard()
        else:
            # Если краткого описания не было, показываем полную клавиатуру
            keyboard = create_analysis_keyboard()
        
        # НЕ УДАЛЯЕМ предыдущие сообщения! Добавляем новое
        await query.message.reply_text(
            "**Выберите тип анализа:**",
            reply_markup=keyboard,
            parse_mode='Markdown'
        )
        return AnalysisStates.ANALYSIS_TYPE_SELECTION.value
    
    elif query.data == "change_analysis_type":
        # ДОБАВЛЯЮ ОТСУТСТВУЮЩИЙ ОБРАБОТЧИК!
        # Проверяем, было ли уже выполнено краткое описание
        if context.user_data.get('summary_done'):
            keyboard = create_other_analysis_keyboard()
        else:
            keyboard = create_analysis_keyboard()
        
        await query.message.reply_text(
            "🎯 **Выберите новый тип анализа:**",
            reply_markup=keyboard,
            parse_mode='Markdown'
        )
        return AnalysisStates.ANALYSIS_TYPE_SELECTION.value
    
    elif query.data == "add_more_images":
        # Ждем загрузки еще одного изображения
        await query.message.reply_text(
            "📷 **Загрузите следующее изображение**\n\n"
            "Отправьте еще одно изображение для добавления к анализу.\n\n"
            "💡 **Совет:** Убедитесь что изображения четкие и текст хорошо виден.",
            parse_mode='Markdown'
        )
        return AnalysisStates.DOCUMENT_UPLOAD.value
    
    elif query.data == "process_single_image":
        # Обрабатываем одно изображение
        logger.info(f"Обработка одного изображения от пользователя {update.effective_user.id}")
        return await process_uploaded_images(update, context, single=True)
    
    elif query.data == "process_all_images":
        # Обрабатываем все загруженные изображения
        logger.info(f"Обработка {len(context.user_data.get('uploaded_images', []))} изображений от пользователя {update.effective_user.id}")
        return await process_uploaded_images(update, context, single=False)
    
    elif query.data == "cancel_analysis":
        # ДОБАВЛЯЮ ОТСУТСТВУЮЩИЙ ОБРАБОТЧИК!
        return await cancel_analysis(update, context)
    
    elif query.data == "back_to_menu":
        # ДОБАВЛЯЮ ОТСУТСТВУЮЩИЙ ОБРАБОТЧИК!
        # Возврат в главное меню через cancel_analysis
        return await cancel_analysis(update, context)
    
    elif query.data == "retry_upload":
        # Повторная попытка загрузки файла
        await query.message.reply_text(
            "🔄 **Повторная загрузка**\n\n"
            "Пожалуйста, загрузите документ заново:\n\n"
            "📋 **Поддерживаемые форматы:**\n"
            "• **Документы:** DOC, DOCX, PDF\n"
            "• **Изображения:** JPG, PNG (сканы)\n"
            "• **Максимальный размер:** 10 МБ",
            parse_mode='Markdown'
        )
        return AnalysisStates.DOCUMENT_UPLOAD.value
    
    elif query.data == "back_to_upload":
        # Возврат к началу загрузки документа
        return await analyze_command(update, context)
    
    # КРИТИЧЕСКИ ВАЖНО: возвращаем состояние для ConversationHandler!
    return AnalysisStates.ANALYSIS_TYPE_SELECTION.value


async def start_text_processing(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Начало извлечения текста из документа
    """
    # Показываем индикатор обработки
    await update.callback_query.message.chat.send_action(ChatAction.TYPING)
    
    file_info = context.user_data['file_info']
    
    # Красивый прогресс-индикатор
    processing_text = (
        f"🔄 **Обработка документа**\n\n"
        f"📄 **Файл:** `{file_info['file_name']}`\n"
        f"📏 **Размер:** {file_info['file_size'] / 1024:.1f} КБ\n"
        f"📝 **Тип:** {'Документ' if file_info['file_type'] == 'document' else 'Изображение (скан)'}\n\n"
        
        "🔍 **Этапы обработки:**\n"
        "▶️ Извлечение текста...\n"
        "⏸️ Анализ содержимого\n"
        "⏸️ Формирование отчета\n\n"
        
        "⏳ *Пожалуйста, подождите...*"
    )
    
    progress_message = await update.callback_query.message.reply_text(
        processing_text,
        parse_mode='Markdown'
    )
    
    # Сохраняем сообщение для обновления прогресса
    context.user_data['progress_message'] = progress_message
    
    try:
        # Извлекаем текст из документа
        extracted_text = await extract_text_from_file(file_info)
        
        if not extracted_text or len(extracted_text.strip()) < 10:
            await update.callback_query.message.reply_text(
                "❌ **Ошибка извлечения текста**\n\n"
                "Не удалось извлечь текст из документа. Возможные причины:\n"
                "• Плохое качество скана\n"
                "• Поврежденный файл\n"
                "• Документ не содержит текста\n\n"
                "Попробуйте загрузить другой файл.",
                parse_mode='Markdown'
            )
            return AnalysisStates.DOCUMENT_UPLOAD.value
        
        # Сохраняем извлеченный текст
        context.user_data['extracted_text'] = extracted_text
        
        # Обновляем прогресс
        progress_message = context.user_data.get('progress_message')
        if progress_message:
            updated_progress = (
                f"🔄 **Обработка документа**\n\n"
                f"📄 **Файл:** `{file_info['file_name']}`\n"
                f"📏 **Размер:** {file_info['file_size'] / 1024:.1f} КБ\n"
                f"📝 **Тип:** {'Документ' if file_info['file_type'] == 'document' else 'Изображение (скан)'}\n\n"
                
                "🔍 **Этапы обработки:**\n"
                "✅ Извлечение текста\n"
                "▶️ Анализ содержимого...\n"
                "⏸️ Формирование отчета\n\n"
                
                "⏳ *Анализируем документ...*"
            )
            
            try:
                await progress_message.edit_text(
                    updated_progress,
                    parse_mode='Markdown'
                )
            except Exception:
                pass  # Игнорируем ошибки редактирования
        
        # Показываем краткую статистику извлеченного текста
        words_count = len(extracted_text.split())
        chars_count = len(extracted_text)
        
        # Определяем качество извлечения
        if chars_count > 1000:
            quality_indicator = "🟢 Отличное качество"
        elif chars_count > 300:
            quality_indicator = "🟡 Хорошее качество"
        else:
            quality_indicator = "🟠 Удовлетворительное качество"
        
        stats_message = (
            f"✅ **Текст успешно извлечен**\n\n"
            f"📊 **Статистика извлечения:**\n"
            f"• **Символов:** {chars_count:,}\n"
            f"• **Слов:** {words_count:,}\n"
            f"• **Качество:** {quality_indicator}\n\n"
            "🧠 Переходим к анализу содержимого..."
        )
        
        await update.callback_query.message.reply_text(
            stats_message,
            parse_mode='Markdown'
        )
        
        # Переходим к анализу
        return await start_analysis_processing(update, context)
        
    except Exception as e:
        logger.error(f"Ошибка извлечения текста: {e}")
        await update.callback_query.message.reply_text(
            "❌ **Техническая ошибка**\n\n"
            "Произошла ошибка при обработке документа.\n"
            "Попробуйте:\n"
            "• Загрузить файл еще раз\n"
            "• Использовать другой формат\n"
            "• Обратиться в поддержку",
            parse_mode='Markdown'
        )
        return ConversationHandler.END


async def start_analysis_processing(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Начало анализа текста через GigaChat
    """
    # Показываем индикатор обработки
    await update.callback_query.message.chat.send_action(ChatAction.TYPING)
    
    analysis_type = context.user_data['analysis_type']
    extracted_text = context.user_data['extracted_text']
    analysis_info = ANALYSIS_TYPES[analysis_type]
    file_info = context.user_data['file_info']
    
    # Обновляем основной прогресс
    progress_message = context.user_data.get('progress_message')
    if progress_message:
        final_progress = (
            f"🔄 **Обработка документа**\n\n"
            f"📄 **Файл:** `{file_info['file_name']}`\n"
            f"📏 **Размер:** {file_info['file_size'] / 1024:.1f} КБ\n"
            f"📝 **Тип:** {'Документ' if file_info['file_type'] == 'document' else 'Изображение (скан)'}\n\n"
            
            "🔍 **Этапы обработки:**\n"
            "✅ Извлечение текста\n"
            "✅ Анализ содержимого\n"
            "▶️ Формирование отчета...\n\n"
            
            "🧠 *Генерируем профессиональный анализ...*"
        )
        
        try:
            await progress_message.edit_text(
                final_progress,
                parse_mode='Markdown'
            )
        except Exception:
            pass
    
    # Детальный статус анализа
    analysis_status = (
        f"🧠 **AI-Анализ документа**\n\n"
        f"{analysis_info['icon']} **{analysis_info['name']}**\n"
        f"📄 **Документ:** `{file_info['file_name']}`\n\n"
        
        "🔍 **Процесс анализа:**\n"
        "▶️ Отправка в GigaChat API...\n"
        "⏸️ Проверка соответствия законодательству\n"
        "⏸️ Выявление проблем и рисков\n"
        "⏸️ Формирование рекомендаций\n\n"
        
        "⏳ *Ожидаемое время: 30-60 секунд*\n"
        "🤖 *Используется профессиональный AI-юрист*"
    )
    
    analysis_message = await update.callback_query.message.reply_text(
        analysis_status,
        parse_mode='Markdown'
    )
    
    # Сохраняем для возможного обновления
    context.user_data['analysis_message'] = analysis_message
    
    try:
        # Получаем имя файла
        file_info = context.user_data['file_info']
        filename = file_info['file_name']
        
        # Анализируем текст через GigaChat
        analysis_result = await analyze_text_with_gigachat(analysis_type, extracted_text, filename)
        
        # Сохраняем результат
        context.user_data['analysis_result'] = analysis_result
        
        # Показываем результат
        return await show_analysis_results(update, context)
        
    except Exception as e:
        logger.error(f"Ошибка анализа: {e}")
        await update.callback_query.message.reply_text(
            "❌ **Ошибка анализа**\n\n"
            "Произошла ошибка при анализе документа.\n"
            "Попробуйте:\n"
            "• Повторить анализ через минуту\n"
            "• Выбрать другой тип анализа\n"
            "• Обратиться в поддержку\n\n"
            "Что делаем дальше?",
            parse_mode='Markdown'
        )
        return AnalysisStates.RESULTS_REVIEW.value


async def show_analysis_results(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Показ результатов анализа с улучшенным форматированием
    """
    analysis_result = context.user_data['analysis_result']
    file_info = context.user_data['file_info']
    analysis_type = context.user_data['analysis_type']
    analysis_info = ANALYSIS_TYPES[analysis_type]
    
    # Завершаем прогресс-индикаторы
    progress_message = context.user_data.get('progress_message')
    analysis_message = context.user_data.get('analysis_message')
    
    if progress_message:
        completion_status = (
            f"✅ **Обработка завершена**\n\n"
            f"📄 **Файл:** `{file_info['file_name']}`\n"
            f"📏 **Размер:** {file_info['file_size'] / 1024:.1f} КБ\n"
            f"📝 **Тип:** {'Документ' if file_info['file_type'] == 'document' else 'Изображение (скан)'}\n\n"
            
            "🔍 **Этапы обработки:**\n"
            "✅ Извлечение текста\n"
            "✅ Анализ содержимого\n"
            "✅ Формирование отчета\n\n"
            
            "🎉 *Анализ успешно завершен!*"
        )
        
        try:
            await progress_message.edit_text(
                completion_status,
                parse_mode='Markdown'
            )
        except Exception:
            pass
    
    if analysis_message:
        try:
            await analysis_message.edit_text(
                f"✅ **Анализ завершен**\n\n"
                f"{analysis_info['icon']} **{analysis_info['name']}**\n"
                f"📄 **Документ:** `{file_info['file_name']}`\n\n"
                "🎯 Результат готов к просмотру ниже ⬇️",
                parse_mode='Markdown'
            )
        except Exception:
            pass
    
    # Создаем расширенную клавиатуру для дополнительных действий
    keyboard = [
        [
            InlineKeyboardButton("🔄 Другой тип анализа", callback_data="change_analysis_type"),
            InlineKeyboardButton("📄 Новый документ", callback_data="upload_new_document")
        ],
        [
            InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu"),
            InlineKeyboardButton("❌ Завершить", callback_data="finish_analysis")
        ]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    # Отправляем результат с улучшенной разбивкой
    await send_long_message(
        update.callback_query.message.chat,
        analysis_result,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return AnalysisStates.RESULTS_REVIEW.value


async def handle_additional_actions(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработчик дополнительных действий
    """
    query = update.callback_query
    await query.answer()
    
    if query.data == "change_analysis_type":
        # Возврат к выбору типа анализа
        keyboard = []
        for analysis_type, info in ANALYSIS_TYPES.items():
            keyboard.append([InlineKeyboardButton(
                info['name'], 
                callback_data=f"analysis_type_{analysis_type}"
            )])
        
        keyboard.append([InlineKeyboardButton("❌ Отмена", callback_data="cancel_analysis")])
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.message.reply_text(
            "🎯 **Выберите новый тип анализа:**",
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
        return AnalysisStates.ANALYSIS_TYPE_SELECTION.value
        
    elif query.data == "upload_new_document":
        # Возврат к загрузке документа
        return await analyze_command(update, context)
        
    elif query.data == "finish_analysis":
        # Завершение анализа
        context.user_data.clear()
        
        # Создаем клавиатуру для выбора дальнейших действий
        keyboard = [
            [InlineKeyboardButton("📊 Новый анализ", callback_data="start_analyze")],
            [InlineKeyboardButton("💬 Консультация", callback_data="menu_consult")],
            [InlineKeyboardButton("📄 Создать документ", callback_data="menu_create")],
            [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.message.reply_text(
            "✅ **Анализ завершен**\n\n"
            "Спасибо за использование AI-юриста!\n\n"
            "Выберите дальнейшие действия:",
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        return ConversationHandler.END
    
    elif query.data == "main_menu":
        # Возврат в главное меню
        context.user_data.clear()
        from bot.handlers import start_command
        return await start_command(update, context)
    
    return AnalysisStates.RESULTS_REVIEW.value


async def process_uploaded_images(update: Update, context: ContextTypes.DEFAULT_TYPE, single: bool = False) -> int:
    """
    Обработка загруженных изображений (одного или нескольких)
    """
    logger.info(f"Начало обработки изображений, single={single}, пользователь {update.effective_user.id}")
    
    uploaded_images = context.user_data.get('uploaded_images', [])
    if not uploaded_images:
        logger.error(f"Нет загруженных изображений для пользователя {update.effective_user.id}")
        await update.callback_query.message.reply_text(
            "❌ **Ошибка:** Нет загруженных изображений",
            parse_mode='Markdown'
        )
        return AnalysisStates.DOCUMENT_UPLOAD.value
    
    # Если single=True, обрабатываем только первое изображение
    images_to_process = [uploaded_images[0]] if single else uploaded_images
    
    logger.info(f"Обрабатываем {len(images_to_process)} изображений")
    
    # Показываем прогресс
    progress_message = await update.callback_query.message.reply_text(
        f"🔄 **Обрабатываю {len(images_to_process)} изображений...**\n\n"
        f"📥 Загружаю файлы...",
        parse_mode='Markdown'
    )
    
    try:
        extracted_texts = []
        
        for i, image_info in enumerate(images_to_process):
            logger.info(f"Обработка изображения {i+1}/{len(images_to_process)}: {image_info['file_name']}")
            
            # Обновляем прогресс (используем безопасное имя файла)
            try:
                await progress_message.edit_text(
                    f"🔄 **Обрабатываю изображения...**\n\n"
                    f"📷 **Изображение {i+1}/{len(images_to_process)}**\n"
                    f"🔍 Распознаю текст...",
                    parse_mode='Markdown'
                )
            except Exception as edit_error:
                logger.warning(f"Не удалось обновить прогресс: {edit_error}")
                # Продолжаем без обновления прогресса
            
            try:
                # Извлекаем текст из изображения
                text = await extract_text_from_file(image_info)
                if text and len(text.strip()) > 3:
                    extracted_texts.append(f"--- Изображение {i+1}: {image_info['file_name']} ---\n{text}")
                    logger.info(f"Текст извлечен из изображения {i+1}: {len(text)} символов")
                else:
                    logger.warning(f"Мало текста извлечено из изображения {i+1}")
                    extracted_texts.append(f"--- Изображение {i+1}: {image_info['file_name']} ---\n⚠️ Текст не распознан или слишком мало")
            
            except Exception as e:
                logger.error(f"Ошибка обработки изображения {i+1}: {e}")
                extracted_texts.append(f"--- Изображение {i+1}: {image_info['file_name']} ---\n❌ Ошибка обработки: {str(e)}")
        
        # Объединяем весь извлеченный текст
        combined_text = "\n\n".join(extracted_texts)
        
        if not combined_text or len(combined_text.strip()) < 10:
            logger.warning(f"Мало текста извлечено из всех изображений: {len(combined_text)} символов")
            
            error_keyboard = [
                [InlineKeyboardButton("🔄 Попробовать еще раз", callback_data="retry_upload")],
                [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
            ]
            
            await progress_message.edit_text(
                "❌ **Не удалось распознать текст**\n\n"
                "Из загруженных изображений извлечено слишком мало текста.\n\n"
                "**Рекомендации:**\n"
                "• Убедитесь что изображения четкие\n"
                "• Текст должен быть хорошо виден\n"
                "• Попробуйте лучшее освещение\n"
                "• Используйте более высокое разрешение",
                reply_markup=InlineKeyboardMarkup(error_keyboard),
                parse_mode='Markdown'
            )
            return AnalysisStates.DOCUMENT_UPLOAD.value
        
        # Сохраняем объединенный текст
        context.user_data['document_text'] = combined_text
        
        # Создаем file_info для совместимости
        total_size = sum(img['file_size'] for img in images_to_process)
        context.user_data['file_info'] = {
            'file_name': f"{len(images_to_process)} изображений",
            'file_size': total_size,
            'file_type': 'image',
            'file_extension': '.jpg'
        }
        
        # Показываем успешный результат
        await progress_message.edit_text(
            f"✅ **Изображения обработаны!**\n\n"
            f"📷 **Обработано:** {len(images_to_process)} изображений\n"
            f"📊 **Общий размер:** {total_size / (1024*1024):.1f} МБ\n"
            f"🔤 **Текст распознан:** {len(combined_text)} символов",
            parse_mode='Markdown'
        )
        
        # Предлагаем выбрать тип анализа
        keyboard = create_analysis_keyboard()
        
        await update.callback_query.message.reply_text(
            f"📄 **Документ получен:** {len(images_to_process)} изображений\n\n"
            f"📊 **Размер:** {total_size / (1024*1024):.1f} МБ\n"
            f"📝 **Текст извлечен:** {len(combined_text)} символов\n\n"
            "**Выберите тип анализа:**",
            reply_markup=keyboard,
            parse_mode='Markdown'
        )
        
        return AnalysisStates.ANALYSIS_TYPE_SELECTION.value
        
    except Exception as e:
        logger.error(f"Критическая ошибка обработки изображений: {e}")
        
        error_keyboard = [
            [InlineKeyboardButton("🔄 Попробовать еще раз", callback_data="retry_upload")],
            [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
        ]
        
        await progress_message.edit_text(
            "❌ **Критическая ошибка**\n\n"
            "Произошла ошибка при обработке изображений.\n"
            "Попробуйте загрузить файлы заново.",
            reply_markup=InlineKeyboardMarkup(error_keyboard),
            parse_mode='Markdown'
        )
        return AnalysisStates.DOCUMENT_UPLOAD.value


async def cancel_analysis(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Отмена анализа документа и возврат в главное меню
    """
    # Очищаем данные
    context.user_data.clear()
    
    # Определяем откуда пришел запрос
    if update.callback_query:
        await update.callback_query.answer()
        
        # Возвращаем в главное меню
        user = update.effective_user
        user_name = user.first_name if user.first_name else "пользователь"
        
        welcome_message = (
            f"🏛️ **AI-Юрист** — ваш персональный правовой помощник\n\n"
            f"👋 Привет, {user_name}!\n\n"
            
            "🎯 **Проверенный сервис:**\n"
            "• 6 отраслей права\n"
            "• 40+ типов документов\n"
            "• Экспорт в Word\n"
            "• Консультаций проведено: 500+\n\n"
            
            "💼 **Выберите нужную услугу:**"
        )
        
        # Создаем клавиатуру главного меню
        keyboard = [
            [InlineKeyboardButton("💬 Юридическая консультация", callback_data="menu_consult")],
            [InlineKeyboardButton("📄 Создание документов", callback_data="menu_create")],
            [InlineKeyboardButton("📊 Анализ документов", callback_data="menu_analyze")],
            [InlineKeyboardButton("ℹ️ Справка и поддержка", callback_data="menu_help")]
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.callback_query.message.reply_text(
            welcome_message,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
    else:
        await update.message.reply_text(
            "❌ Анализ документа отменен.\n"
            "Используйте кнопки ниже для выбора действий."
        )
    
    return ConversationHandler.END


# Вспомогательные функции

async def extract_text_from_file(file_info: Dict[str, Any]) -> str:
    """
    Извлечение текста из файла
    """
    file_extension = file_info['file_extension'].lower()
    file_type = file_info['file_type']
    
    try:
        # Скачиваем файл во временную директорию
        temp_file_path = await download_telegram_file(file_info['file_id'], file_extension)
        
        extracted_text = ""
        
        if file_type == 'document':
            if file_extension in ['.doc', '.docx']:
                extracted_text = await extract_text_from_docx(temp_file_path)
            elif file_extension == '.pdf':
                extracted_text = await extract_text_from_pdf(temp_file_path)
        
        elif file_type == 'image':
            extracted_text = await extract_text_from_image(temp_file_path)
        
        # Удаляем временный файл
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        
        # Очищаем и нормализуем текст
        cleaned_text = clean_extracted_text(extracted_text)
        
        return cleaned_text
        
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из файла {file_info['file_name']}: {e}")
        raise Exception(f"Не удалось извлечь текст из файла: {str(e)}")


async def analyze_text_with_gigachat(analysis_type: str, text: str, filename: str = "документ") -> str:
    """
    Анализ текста через GigaChat API с улучшенным логированием
    """
    logger.info(f"Начинаем анализ документа '{filename}' типа '{analysis_type}'")
    logger.info(f"Размер текста для анализа: {len(text)} символов")
    
    try:
        # Импортируем клиент GigaChat
        from ai_gigachat.client import gigachat_client
        
        # Проверяем размер текста
        if len(text) > 100000:  # 100KB текста
            logger.warning(f"Большой объем текста для анализа: {len(text)} символов")
        
        # Засекаем время начала анализа
        import time
        start_time = time.time()
        
        # Выполняем анализ через GigaChat API
        logger.info("Отправляем запрос в GigaChat API...")
        
        analysis_result = await gigachat_client.analyze_document(
            document_text=text,
            analysis_type=analysis_type,
            filename=filename
        )
        
        # Засекаем время завершения
        end_time = time.time()
        duration = end_time - start_time
        
        logger.info(f"Анализ документа '{filename}' успешно завершен за {duration:.2f} секунд")
        logger.info(f"Размер результата анализа: {len(analysis_result)} символов")
        
        return analysis_result
        
    except Exception as e:
        logger.error(f"ОШИБКА при анализе документа '{filename}':")
        logger.error(f"  Тип ошибки: {type(e).__name__}")
        logger.error(f"  Сообщение: {str(e)}")
        logger.error(f"  Тип анализа: {analysis_type}")
        logger.error(f"  Размер текста: {len(text)} символов")
        
        # Проверяем тип ошибки для более точного ответа
        error_type = type(e).__name__
        error_message = str(e)
        
        # Возвращаем резервный ответ при ошибке
        analysis_info = ANALYSIS_TYPES[analysis_type]
        
        if "timeout" in error_message.lower() or "time" in error_message.lower():
            error_description = "Превышено время ожидания ответа от AI-сервиса"
            suggestions = [
                "• Попробуйте повторить анализ через 1-2 минуты",
                "• Уменьшите размер документа если возможно",
                "• Разбейте большой документ на части"
            ]
        elif "connection" in error_message.lower() or "network" in error_message.lower():
            error_description = "Проблемы с сетевым соединением"
            suggestions = [
                "• Проверьте подключение к интернету",
                "• Попробуйте повторить запрос через минуту",
                "• Обратитесь в поддержку если проблема повторяется"
            ]
        elif "rate" in error_message.lower() or "limit" in error_message.lower():
            error_description = "Превышен лимит запросов к AI-сервису"
            suggestions = [
                "• Подождите несколько минут перед повтором",
                "• Попробуйте использовать меньший документ",
                "• Обратитесь в поддержку для увеличения лимитов"
            ]
        else:
            error_description = "Временные проблемы с AI-сервисом"
            suggestions = [
                "• Попробуйте повторить анализ через несколько минут",
                "• Убедитесь, что документ содержит читаемый текст",
                "• Обратитесь в поддержку, если проблема повторяется"
            ]
        
        return f"""❌ **Ошибка анализа документа**

{analysis_info['icon']} **{analysis_info['name']}**
📄 **Файл:** {filename}

**Причина:** {error_description}

**Что можно сделать:**
{chr(10).join(suggestions)}

🔄 Используйте кнопки ниже для дальнейших действий"""


async def send_long_message(chat, text: str, **kwargs) -> None:
    """
    Умная отправка длинного сообщения с разбивкой на части по смыслу
    """
    MAX_MESSAGE_LENGTH = 4096
    
    if len(text) <= MAX_MESSAGE_LENGTH:
        await chat.send_message(text, **kwargs)
        return
    
    # Разбиваем текст на смысловые части
    parts = smart_split_message(text, MAX_MESSAGE_LENGTH)
    
    # Отправляем части по очереди
    for i, part in enumerate(parts):
        if i == len(parts) - 1:
            # Последняя часть - с клавиатурой
            await chat.send_message(part, **kwargs)
        else:
            # Промежуточные части - без клавиатуры и с индикатором продолжения
            kwargs_copy = kwargs.copy()
            kwargs_copy.pop('reply_markup', None)
            
            # Добавляем красивый индикатор продолжения
            continuation_indicator = f"\n\n📄 **Часть {i+1}/{len(parts)}** • *Продолжение следует...*"
            await chat.send_message(part + continuation_indicator, **kwargs_copy)
            
            # Небольшая пауза между сообщениями для лучшего восприятия
            await asyncio.sleep(0.5)


def smart_split_message(text: str, max_length: int) -> List[str]:
    """
    Умная разбивка сообщения по смыслу с сохранением форматирования
    """
    if len(text) <= max_length:
        return [text]
    
    parts = []
    current_text = text
    
    while len(current_text) > max_length:
        # Приоритеты для разбивки (от наиболее предпочтительного к наименее)
        split_patterns = [
            '\n\n**',          # Новый раздел с заголовком
            '\n**',            # Заголовок
            '\n\n',            # Двойной перенос строки
            '\n•',             # Список
            '\n-',             # Список с тире
            '\n',              # Обычный перенос строки
            '. ',              # Конец предложения
            ', ',              # Запятая
            ' '                # Пробел
        ]
        
        split_index = -1
        
        # Ищем лучшее место для разбивки
        for pattern in split_patterns:
            # Ищем последнее вхождение паттерна в допустимом диапазоне
            temp_index = current_text.rfind(pattern, 0, max_length - 100)  # Оставляем буфер
            if temp_index > max_length // 2:  # Не разбиваем слишком рано
                split_index = temp_index + len(pattern)
                break
        
        # Если не нашли подходящее место, разбиваем по максимальной длине
        if split_index == -1:
            split_index = max_length
        
        # Добавляем часть
        part = current_text[:split_index].strip()
        if part:
            parts.append(part)
        
        # Переходим к оставшемуся тексту
        current_text = current_text[split_index:].strip()
    
    # Добавляем оставшийся текст
    if current_text:
        parts.append(current_text)
    
    return parts


# Функции для извлечения текста из разных типов файлов

async def download_telegram_file(file_id: str, file_extension: str) -> str:
    """
    Скачивание файла из Telegram во временную директорию
    """
    from telegram import Bot
    from config import TELEGRAM_TOKEN
    
    bot = Bot(token=TELEGRAM_TOKEN)
    
    # Получаем информацию о файле
    file = await bot.get_file(file_id)
    
    # Создаем временный файл
    temp_dir = tempfile.gettempdir()
    temp_file_path = os.path.join(temp_dir, f"telegram_file_{file_id}{file_extension}")
    
    # Скачиваем файл
    await file.download_to_drive(temp_file_path)
    
    return temp_file_path


async def extract_text_from_docx(file_path: str) -> str:
    """
    Извлечение текста из DOCX файла
    """
    try:
        # Загружаем документ
        doc = DocxDocument(file_path)
        
        # Извлекаем текст из всех параграфов
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        extracted_text = '\n'.join(text_parts)
        
        if not extracted_text.strip():
            raise Exception("Документ не содержит текста или поврежден")
        
        return extracted_text
        
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из DOCX: {e}")
        raise Exception(f"Не удалось обработать DOCX файл: {str(e)}")


async def extract_text_from_pdf(file_path: str) -> str:
    """
    Извлечение текста из PDF файла с таймаутами и улучшенным логированием
    """
    logger.info(f"Начало извлечения текста из PDF: {file_path}")
    
    extracted_text = ""
    
    try:
        # Проверяем размер файла
        file_size = os.path.getsize(file_path)
        logger.info(f"Размер PDF файла: {file_size} байт ({file_size / (1024*1024):.1f} МБ)")
        
        # Устанавливаем таймаут в зависимости от размера файла
        timeout_seconds = min(120, max(30, file_size // (1024 * 1024) * 10))  # 10 сек на МБ, мин 30, макс 120
        logger.info(f"Установлен таймаут: {timeout_seconds} секунд")
        
        # Сначала пытаемся извлечь текст напрямую через PyMuPDF (быстрее)
        logger.info("Попытка извлечения текста через PyMuPDF...")
        
        async def extract_with_pymupdf():
            pdf_document = fitz.open(file_path)
            logger.info(f"PDF открыт, количество страниц: {pdf_document.page_count}")
            
            text_parts = []
            for page_num in range(pdf_document.page_count):
                logger.debug(f"Обработка страницы {page_num + 1}/{pdf_document.page_count}")
                # Правильный API для PyMuPDF - используем индексацию
                page = pdf_document[page_num]
                page_text = page.get_text()
                
                if page_text.strip():
                    text_parts.append(f"\n--- Страница {page_num + 1} ---\n{page_text}\n")
                    logger.debug(f"Страница {page_num + 1}: извлечено {len(page_text)} символов")
                else:
                    logger.warning(f"Страница {page_num + 1}: текст не найден")
            
            pdf_document.close()
            return "\n".join(text_parts)
        
        # Выполняем с таймаутом
        try:
            extracted_text = await asyncio.wait_for(extract_with_pymupdf(), timeout=timeout_seconds)
            logger.info(f"PyMuPDF успешно завершен: {len(extracted_text)} символов")
        except asyncio.TimeoutError:
            logger.error(f"Таймаут PyMuPDF ({timeout_seconds}s), пробуем альтернативный метод")
            raise Exception(f"Обработка PDF заняла слишком много времени (>{timeout_seconds}s)")
        
        # Если текста достаточно, возвращаем результат
        if len(extracted_text.strip()) > 50:
            logger.info(f"Достаточно текста извлечено: {len(extracted_text)} символов")
            return extracted_text
        
        # Если текста мало, пробуем OCR
        logger.warning(f"Мало текста в PDF ({len(extracted_text)} символов), пытаемся OCR...")
        try:
            ocr_text = await asyncio.wait_for(
                extract_text_from_pdf_with_ocr(file_path), 
                timeout=timeout_seconds * 2  # OCR может занять больше времени
            )
            logger.info(f"OCR завершен: {len(ocr_text) if ocr_text else 0} символов")
            return ocr_text if ocr_text else extracted_text
        except asyncio.TimeoutError:
            logger.error(f"Таймаут OCR ({timeout_seconds * 2}s)")
            return extracted_text  # Возвращаем что есть
        
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из PDF: {e}")
        
        # В случае ошибки пробуем альтернативный метод через PyPDF2
        try:
            logger.info("Пробуем альтернативный метод PyPDF2...")
            alt_text = await asyncio.wait_for(
                extract_text_from_pdf_pypdf2(file_path), 
                timeout=60
            )
            logger.info(f"PyPDF2 успешно завершен: {len(alt_text)} символов")
            return alt_text
        except asyncio.TimeoutError:
            logger.error("Таймаут PyPDF2 (60s)")
            raise Exception("PDF обработка заняла слишком много времени")
        except Exception as e2:
            logger.error(f"Ошибка альтернативного извлечения из PDF: {e2}")
            raise Exception(f"Не удалось обработать PDF файл: {str(e)}")


async def extract_text_from_pdf_pypdf2(file_path: str) -> str:
    """
    Альтернативное извлечение текста из PDF через PyPDF2
    """
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                if page_text.strip():
                    text_parts.append(f"--- Страница {page_num + 1} ---\n{page_text}")
            
            extracted_text = '\n'.join(text_parts)
            
            if not extracted_text.strip():
                raise Exception("PDF не содержит извлекаемого текста")
            
            return extracted_text
            
    except Exception as e:
        raise Exception(f"Ошибка PyPDF2: {str(e)}")


async def extract_text_from_pdf_with_ocr(file_path: str) -> str:
    """
    Извлечение текста из PDF с помощью OCR (для сканированных документов)
    """
    try:
        # Проверяем доступность Tesseract
        if not is_tesseract_available():
            raise Exception("OCR недоступен: требуется установка Tesseract OCR")
        
        # Конвертируем PDF в изображения
        from pdf2image import convert_from_path
        
        # Конвертируем страницы PDF в изображения
        images = convert_from_path(file_path, dpi=300, fmt='jpeg')
        
        text_parts = []
        for i, image in enumerate(images):
            # Применяем OCR к каждой странице
            page_text = pytesseract.image_to_string(image, lang='rus+eng')
            
            if page_text.strip():
                text_parts.append(f"--- Страница {i + 1} ---\n{page_text}")
        
        extracted_text = '\n'.join(text_parts)
        
        if not extracted_text.strip():
            raise Exception("OCR не смог распознать текст в PDF")
        
        return extracted_text
        
    except Exception as e:
        logger.error(f"Ошибка OCR для PDF: {e}")
        raise Exception(f"Не удалось распознать текст в сканированном PDF: {str(e)}")


async def extract_text_from_image(file_path: str) -> str:
    """
    Извлечение текста из изображения с помощью OCR
    """
    try:
        logger.info(f"Начинаем OCR обработку изображения: {file_path}")
        
        # Проверяем доступность Tesseract
        if not is_tesseract_available():
            logger.error("Tesseract OCR недоступен")
            raise Exception("OCR недоступен: требуется установка Tesseract OCR.\n"
                          "Для анализа изображений необходимо установить Tesseract OCR в системе.")
        
        # Проверяем существование файла
        if not os.path.exists(file_path):
            logger.error(f"Файл изображения не найден: {file_path}")
            raise Exception("Файл изображения не найден")
        
        logger.info(f"Размер файла: {os.path.getsize(file_path)} байт")
        
        # Открываем изображение
        image = Image.open(file_path)
        logger.info(f"Изображение открыто: {image.size}, режим: {image.mode}")
        
        # Конвертируем в RGB если нужно
        if image.mode != 'RGB':
            image = image.convert('RGB')
            logger.info("Изображение конвертировано в RGB")
        
        # Улучшаем качество изображения для лучшего OCR
        # Увеличиваем размер если изображение маленькое
        width, height = image.size
        if width < 1000 or height < 1000:
            scale_factor = max(1000 / width, 1000 / height)
            new_width = int(width * scale_factor)
            new_height = int(height * scale_factor)
            image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            logger.info(f"Изображение увеличено до {new_width}x{new_height}")
        
        # Применяем OCR с разными настройками
        logger.info("Применяем OCR...")
        
        # Пробуем несколько конфигураций OCR
        ocr_configs = [
            '--psm 3',  # Полностью автоматическая сегментация страницы (по умолчанию)
            '--psm 6',  # Единый блок текста
            '--psm 4',  # Одна колонка текста различных размеров
        ]
        
        extracted_text = ""
        for config in ocr_configs:
            try:
                text = pytesseract.image_to_string(image, lang='rus+eng', config=config)
                if text and len(text.strip()) > len(extracted_text.strip()):
                    extracted_text = text
                    logger.info(f"OCR успешен с конфигурацией {config}, длина текста: {len(text.strip())}")
            except Exception as e:
                logger.warning(f"OCR с конфигурацией {config} не удался: {e}")
                continue
        
        if not extracted_text or len(extracted_text.strip()) < 5:
            logger.warning(f"OCR вернул мало текста: '{extracted_text[:100]}...'")
            
            # Попробуем только английский язык
            try:
                text_eng = pytesseract.image_to_string(image, lang='eng')
                if len(text_eng.strip()) > len(extracted_text.strip()):
                    extracted_text = text_eng
                    logger.info(f"OCR с английским языком дал лучший результат: {len(text_eng.strip())}")
            except Exception as e:
                logger.warning(f"OCR только с английским не удался: {e}")
            
            # Если все еще мало текста, возвращаем что есть с предупреждением
            if len(extracted_text.strip()) < 5:
                return f"⚠️ OCR распознал очень мало текста.\n\nВозможные причины:\n• Плохое качество изображения\n• Неразборчивый текст\n• Необычный шрифт\n\nРаспознанный текст:\n{extracted_text.strip()}"
        
        logger.info(f"OCR завершен успешно, итоговая длина: {len(extracted_text.strip())}")
        return extracted_text
        
    except Exception as e:
        logger.error(f"Критическая ошибка OCR для изображения {file_path}: {e}")
        return f"❌ Ошибка OCR: {str(e)}\n\nПопробуйте:\n• Загрузить изображение лучшего качества\n• Использовать более контрастное изображение\n• Убедиться что текст четко виден"


def is_tesseract_available() -> bool:
    """
    Проверка доступности Tesseract OCR
    """
    try:
        import subprocess
        result = subprocess.run(['tesseract', '--version'], 
                              capture_output=True, text=True, timeout=10)
        return result.returncode == 0
    except (subprocess.TimeoutExpired, subprocess.CalledProcessError, FileNotFoundError):
        return False


def escape_markdown(text: str) -> str:
    """
    Экранирует специальные символы Markdown
    """
    if not text:
        return ""
    
    # Экранируем основные символы Markdown
    escape_chars = ['_', '*', '[', ']', '(', ')', '~', '`', '>', '#', '+', '-', '=', '|', '{', '}', '.', '!']
    
    for char in escape_chars:
        text = text.replace(char, f'\\{char}')
    
    return text


def clean_extracted_text(text: str) -> str:
    """
    Очистка и нормализация извлеченного текста
    """
    if not text:
        return ""
    
    # Заменяем множественные пробелы на одинарные (но сохраняем переносы строк)
    cleaned = re.sub(r'[ \t]+', ' ', text)
    
    # Удаляем повторяющиеся переносы строк (больше 2 подряд)
    cleaned = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned)
    
    # Убираем пробелы в начале и конце строк
    lines = [line.strip() for line in cleaned.split('\n')]
    cleaned = '\n'.join(lines)
    
    # Убираем пробелы в начале и конце всего текста
    cleaned = cleaned.strip()
    
    # Для OCR результатов не требуем минимальную длину
    # Возвращаем текст как есть, даже если он короткий
    return cleaned

def create_analysis_keyboard():
    """Создает клавиатуру для выбора типа анализа"""
    keyboard = []
    
    # Добавляем кнопку "Краткое описание" первой
    keyboard.append([InlineKeyboardButton(
        ANALYSIS_TYPES['document_summary']['name'], 
        callback_data="analysis_type_document_summary"
    )])
    
    # Добавляем остальные типы анализа
    for analysis_type, info in ANALYSIS_TYPES.items():
        if analysis_type != 'document_summary':  # Пропускаем, так как уже добавили
            keyboard.append([InlineKeyboardButton(
                info['name'], 
                callback_data=f"analysis_type_{analysis_type}"
            )])
    
    # Кнопки управления
    keyboard.extend([
        [InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")],
        [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
    ])
    
    return InlineKeyboardMarkup(keyboard) 

def create_other_analysis_keyboard():
    """Создает клавиатуру с остальными типами анализа (без краткого описания)"""
    keyboard = []
    
    # Добавляем все типы анализа кроме document_summary
    for analysis_type, info in ANALYSIS_TYPES.items():
        if analysis_type != 'document_summary':
            keyboard.append([InlineKeyboardButton(
                info['name'], 
                callback_data=f"analysis_type_{analysis_type}"
            )])
    
    # Кнопки управления
    keyboard.extend([
        [InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")],
        [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
    ])
    
    return InlineKeyboardMarkup(keyboard) 