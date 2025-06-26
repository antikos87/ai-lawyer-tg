#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Обработчики для создания юридических документов
"""

import logging
from enum import Enum
from typing import Dict, Any, List, Optional
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes, ConversationHandler  # type: ignore
from telegram.constants import ChatAction  # type: ignore

from ai_gigachat.client import gigachat_client
from .analysis_handlers import send_long_message, smart_split_message
import config

logger = logging.getLogger(__name__)


class DocumentStates(Enum):
    """Состояния диалога создания документов"""
    DOCUMENT_TYPE_SELECTION = 0
    DOCUMENT_SUBTYPE_SELECTION = 1
    DATA_COLLECTION = 2
    DOCUMENT_GENERATION = 3
    DOCUMENT_REVIEW = 4
    DOCUMENT_EDITING = 5
    DOCUMENT_FINALIZATION = 6


# Типы документов и их подтипы
DOCUMENT_TYPES = {
    "contract": {
        "name": "📄 Договор",
        "subtypes": {
            "rent": "Аренды",
            "sale": "Купли-продажи", 
            "service": "Оказания услуг",
            "work": "Выполнения работ",
            "employment": "Трудовой",
            "loan": "Займа"
        }
    },
    "lawsuit": {
        "name": "⚖️ Исковое заявление",
        "subtypes": {
            "property": "Имущественный спор",
            "family": "Семейный спор",
            "labor": "Трудовой спор",
            "compensation": "Возмещение ущерба",
            "debt": "Взыскание долга"
        }
    },
    "claim": {
        "name": "📧 Досудебная претензия",
        "subtypes": {
            "payment": "О взыскании долга",
            "quality": "О некачественном товаре/услуге",
            "contract": "О нарушении договора",
            "refund": "О возврате денежных средств"
        }
    },
    "agreement": {
        "name": "🤝 Соглашение",
        "subtypes": {
            "settlement": "Мировое соглашение",
            "alimony": "Об алиментах",
            "property": "О разделе имущества",
            "cooperation": "О сотрудничестве"
        }
    },
    "power_of_attorney": {
        "name": "📋 Доверенность",
        "subtypes": {
            "general": "Генеральная",
            "property": "На недвижимость",
            "vehicle": "На автомобиль",
            "bank": "В банк",
            "court": "В суд"
        }
    },
    "application": {
        "name": "📝 Заявление",
        "subtypes": {
            "court": "В суд",
            "police": "В полицию",
            "administration": "В администрацию",
            "employer": "Работодателю"
        }
    },
    "protocol": {
        "name": "📊 Протокол",
        "subtypes": {
            "meeting": "Собрания",
            "inspection": "Осмотра",
            "violation": "Нарушения",
            "handover": "Передачи"
        }
    },
    "act": {
        "name": "📑 Акт",
        "subtypes": {
            "acceptance": "Приема-передачи",
            "inspection": "Осмотра",
            "completion": "Выполненных работ",
            "damage": "О повреждении"
        }
    }
}


async def create_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработчик команды /create
    Запускает процесс создания юридического документа
    """
    user = update.effective_user
    user_name = user.first_name if user.first_name else "пользователь"
    
    # Очищаем предыдущие данные документа
    context.user_data.pop('document_data', None)
    context.user_data.pop('document_type', None)
    context.user_data.pop('document_subtype', None)
    context.user_data.pop('generated_document', None)
    
    welcome_text = (
        f"📄 Привет, {user_name}!\n\n"
        "🎯 Я помогу вам создать юридический документ.\n\n"
        "💡 **Доступные типы документов:**\n"
        "• Договоры различных типов\n"
        "• Исковые заявления\n"
        "• Досудебные претензии\n"
        "• Соглашения и доверенности\n"
        "• Заявления и протоколы\n\n"
        "Выберите тип документа:"
    )
    
    # Создаем клавиатуру с типами документов
    keyboard = []
    for doc_type, info in DOCUMENT_TYPES.items():
        keyboard.append([InlineKeyboardButton(info["name"], callback_data=f"doctype_{doc_type}")])
    
    # Добавляем кнопку главного меню
    keyboard.append([InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    # Проверяем, вызвано ли из callback query или обычного message
    if update.callback_query:
        # Вызвано из callback query (например, из главного меню)
        await update.callback_query.answer()
        await update.callback_query.message.reply_text(
            welcome_text,
            reply_markup=reply_markup
        )
    else:
        # Вызвано из обычного сообщения (команда /create)
        await update.message.reply_text(
            welcome_text,
            reply_markup=reply_markup
        )
    
    return DocumentStates.DOCUMENT_TYPE_SELECTION.value


async def document_type_selected(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработчик выбора типа документа
    """
    query = update.callback_query
    await query.answer()
    
    doc_type = query.data.replace("doctype_", "")
    context.user_data['document_type'] = doc_type
    
    doc_info = DOCUMENT_TYPES[doc_type]
    
    subtype_text = (
        f"📄 **Выбран тип: {doc_info['name']}**\n\n"
        "Теперь выберите подтип документа:"
    )
    
    # Создаем клавиатуру с подтипами
    keyboard = []
    for subtype_key, subtype_name in doc_info["subtypes"].items():
        callback_data = f"docsubtype_{subtype_key}"
        keyboard.append([InlineKeyboardButton(f"📋 {subtype_name}", callback_data=callback_data)])
    
    # Кнопки навигации
    keyboard.extend([
        [InlineKeyboardButton("⬅️ Назад к типам", callback_data="back_to_types")],
        [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
    ])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        subtype_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return DocumentStates.DOCUMENT_SUBTYPE_SELECTION.value


async def document_subtype_selected(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработчик выбора подтипа документа
    Начинает сбор данных
    """
    query = update.callback_query
    await query.answer()
    
    subtype = query.data.replace("docsubtype_", "")
    context.user_data['document_subtype'] = subtype
    
    # Инициализируем данные документа
    context.user_data['document_data'] = {}
    context.user_data['current_question'] = 0
    
    doc_type = context.user_data['document_type']
    doc_info = DOCUMENT_TYPES[doc_type]
    subtype_name = doc_info["subtypes"][subtype]
    
    confirmation_text = (
        f"✅ **Выбран документ: {doc_info['name']} - {subtype_name}**\n\n"
        "📝 Теперь я задам вам несколько вопросов для заполнения документа.\n"
        "Отвечайте максимально подробно для качественного результата.\n\n"
        "Готовы начать?"
    )
    
    keyboard = [
        [InlineKeyboardButton("▶️ Начать заполнение", callback_data="start_data_collection")],
        [InlineKeyboardButton("⬅️ Назад к подтипам", callback_data="back_to_subtypes")],
        [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        confirmation_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return DocumentStates.DATA_COLLECTION.value


async def start_data_collection(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Начало сбора данных для документа
    """
    query = update.callback_query
    await query.answer()
    
    # Переходим к первому вопросу
    return await ask_next_question(update, context, is_callback=True)


async def ask_next_question(update: Update, context: ContextTypes.DEFAULT_TYPE, is_callback: bool = False) -> int:
    """
    Задает следующий вопрос для сбора данных
    """
    doc_type = context.user_data['document_type']
    doc_subtype = context.user_data['document_subtype']
    current_question = context.user_data.get('current_question', 0)
    
    # Получаем список вопросов для данного типа документа
    questions = get_questions_for_document(doc_type, doc_subtype)
    
    if current_question >= len(questions):
        # Все вопросы заданы, переходим к генерации
        return await start_document_generation(update, context, is_callback)
    
    question_data = questions[current_question]
    question_text = (
        f"❓ **Вопрос {current_question + 1} из {len(questions)}**\n\n"
        f"{question_data['question']}\n\n"
        f"💡 *{question_data['hint']}*"
    )
    
    keyboard = []
    
    # Если есть варианты ответов, добавляем их как кнопки
    if 'options' in question_data:
        for option in question_data['options']:
            keyboard.append([InlineKeyboardButton(option, callback_data=f"answer_{option}")])
        keyboard.append([InlineKeyboardButton("✍️ Ввести свой вариант", callback_data="custom_answer")])
    
    # Кнопки навигации
    if current_question > 0:
        keyboard.append([InlineKeyboardButton("⬅️ Предыдущий вопрос", callback_data="prev_question")])
    
    keyboard.extend([
        [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
    ])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    if is_callback:
        await update.callback_query.message.reply_text(
            question_text,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
    else:
        await update.message.reply_text(
            question_text,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
    
    return DocumentStates.DATA_COLLECTION.value


def get_questions_for_document(doc_type: str, doc_subtype: str) -> List[Dict[str, Any]]:
    """
    Возвращает список вопросов для конкретного типа документа
    """
    # Базовые вопросы для разных типов документов
    questions_map = {
        "contract": {
            "rent": [
                {
                    "key": "landlord_name",
                    "question": "Укажите ФИО арендодателя (полностью)",
                    "hint": "Например: Иванов Иван Иванович"
                },
                {
                    "key": "tenant_name", 
                    "question": "Укажите ФИО арендатора (полностью)",
                    "hint": "Например: Петров Петр Петрович"
                },
                {
                    "key": "property_address",
                    "question": "Укажите полный адрес арендуемого имущества",
                    "hint": "Включите город, улицу, дом, квартиру"
                },
                {
                    "key": "property_type",
                    "question": "Что сдается в аренду?",
                    "hint": "Выберите тип имущества",
                    "options": ["Квартира", "Комната", "Дом", "Офис", "Торговое помещение", "Склад"]
                },
                {
                    "key": "rent_amount",
                    "question": "Укажите размер арендной платы (в рублях)",
                    "hint": "Например: 50000"
                },
                {
                    "key": "rent_period",
                    "question": "Как часто вносится арендная плата?",
                    "hint": "Выберите периодичность",
                    "options": ["Ежемесячно", "Ежеквартально", "Раз в полгода", "Ежегодно"]
                },
                {
                    "key": "contract_duration",
                    "question": "На какой срок заключается договор?",
                    "hint": "Например: 1 год, 6 месяцев, 2 года"
                }
            ],
            "sale": [
                {
                    "key": "seller_name",
                    "question": "Укажите ФИО продавца (полностью)",
                    "hint": "Например: Иванов Иван Иванович"
                },
                {
                    "key": "buyer_name",
                    "question": "Укажите ФИО покупателя (полностью)", 
                    "hint": "Например: Петров Петр Петрович"
                },
                {
                    "key": "item_description",
                    "question": "Что продается? Опишите предмет сделки",
                    "hint": "Укажите наименование, характеристики, состояние"
                },
                {
                    "key": "sale_price",
                    "question": "Укажите цену продажи (в рублях)",
                    "hint": "Например: 1500000"
                },
                {
                    "key": "payment_method",
                    "question": "Как будет происходить оплата?",
                    "hint": "Выберите способ расчета",
                    "options": ["Наличные", "Безналичный расчет", "Рассрочка", "Ипотека"]
                }
            ]
        },
        "lawsuit": {
            "debt": [
                {
                    "key": "plaintiff_name",
                    "question": "Укажите ФИО истца (полностью)",
                    "hint": "Ваши данные или данные представляемого лица"
                },
                {
                    "key": "defendant_name", 
                    "question": "Укажите ФИО ответчика (полностью)",
                    "hint": "Лицо, с которого взыскивается долг"
                },
                {
                    "key": "debt_amount",
                    "question": "Укажите размер долга (в рублях)",
                    "hint": "Основная сумма без процентов и штрафов"
                },
                {
                    "key": "debt_basis",
                    "question": "На основании чего возник долг?",
                    "hint": "Договор, расписка, решение суда и т.д."
                },
                {
                    "key": "debt_date",
                    "question": "Когда должен был быть погашен долг?",
                    "hint": "Дата, когда ответчик должен был вернуть деньги"
                }
            ]
        },
        "claim": {
            "payment": [
                {
                    "key": "creditor_name",
                    "question": "Укажите ваше ФИО (кредитор)",
                    "hint": "Тот, кому должны деньги"
                },
                {
                    "key": "debtor_name",
                    "question": "Укажите ФИО должника",
                    "hint": "Тот, кто должен деньги"
                },
                {
                    "key": "debt_amount",
                    "question": "Размер задолженности (в рублях)",
                    "hint": "Сумма основного долга"
                },
                {
                    "key": "debt_reason",
                    "question": "По какому основанию возникла задолженность?",
                    "hint": "Договор, услуга, товар и т.д."
                }
            ]
        }
    }
    
    # Возвращаем вопросы для конкретного типа или базовые вопросы
    if doc_type in questions_map and doc_subtype in questions_map[doc_type]:
        return questions_map[doc_type][doc_subtype]
    
    # Базовые вопросы, если специфичных нет
    return [
        {
            "key": "party1_name",
            "question": "Укажите ФИО первой стороны",
            "hint": "Полные данные первой стороны документа"
        },
        {
            "key": "party2_name", 
            "question": "Укажите ФИО второй стороны",
            "hint": "Полные данные второй стороны документа"
        },
        {
            "key": "subject",
            "question": "Опишите предмет/суть документа",
            "hint": "О чем этот документ, что регулирует"
        }
    ]


# Функции навигации и вспомогательные обработчики
async def back_to_types(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Возврат к выбору типа документа"""
    query = update.callback_query
    await query.answer()
    
    # Очищаем текущий выбор
    context.user_data.pop('document_type', None)
    context.user_data.pop('document_subtype', None)
    
    # Показываем выбор типов документов
    keyboard = []
    for doc_type, doc_info in DOCUMENT_TYPES.items():
        keyboard.append([InlineKeyboardButton(doc_info["name"], callback_data=f"doctype_{doc_type}")])
    
    keyboard.append([InlineKeyboardButton("❌ Отмена", callback_data="cancel_document")])
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        "📋 **Выберите тип документа для создания:**\n\n"
        "Каждый тип документа имеет свои особенности и требования. "
        "Выберите подходящий вариант:",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return DocumentStates.DOCUMENT_TYPE_SELECTION.value


async def back_to_subtypes(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Возврат к выбору подтипа документа"""
    query = update.callback_query
    await query.answer()
    
    doc_type = context.user_data['document_type']
    doc_info = DOCUMENT_TYPES[doc_type]
    
    # Очищаем выбранный подтип
    context.user_data.pop('document_subtype', None)
    
    # Показываем подтипы для выбранного типа
    keyboard = []
    for subtype_key, subtype_name in doc_info["subtypes"].items():
        keyboard.append([InlineKeyboardButton(subtype_name, callback_data=f"docsubtype_{subtype_key}")])
    
    keyboard.extend([
        [InlineKeyboardButton("⬅️ Назад к типам", callback_data="back_to_types")],
        [InlineKeyboardButton("❌ Отмена", callback_data="cancel_document")]
    ])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        f"📄 **{doc_info['name']}**\n\n"
        f"{doc_info['description']}\n\n"
        "**Выберите подтип документа:**",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return DocumentStates.DOCUMENT_SUBTYPE_SELECTION.value


async def cancel_document_creation(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Отмена создания документа"""
    # Очищаем данные
    context.user_data.pop('document_data', None)
    context.user_data.pop('document_type', None)
    context.user_data.pop('document_subtype', None)
    context.user_data.pop('generated_document', None)
    
    # Определяем откуда пришел запрос - callback или сообщение
    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.reply_text(
            "❌ Создание документа отменено.\n"
            "Для создания нового документа используйте /create"
        )
    else:
        await update.message.reply_text(
            "❌ Создание документа отменено.\n"
            "Для создания нового документа используйте /create"
        )
    return ConversationHandler.END


async def process_answer(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработка ответа пользователя на вопрос
    """
    if update.callback_query:
        # Ответ через кнопку
        query = update.callback_query
        await query.answer()
        
        if query.data.startswith("answer_"):
            # Выбран готовый вариант ответа
            answer = query.data.replace("answer_", "")
            return await save_answer_and_continue(update, context, answer, is_callback=True)
        elif query.data == "custom_answer":
            # Пользователь хочет ввести свой ответ
            await query.message.reply_text(
                "✍️ Введите ваш ответ:",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("❌ Отмена", callback_data="cancel_custom")]
                ])
            )
            return DocumentStates.DATA_COLLECTION.value
        elif query.data == "prev_question":
            # Переход к предыдущему вопросу
            context.user_data['current_question'] = max(0, context.user_data.get('current_question', 0) - 1)
            return await ask_next_question(update, context, is_callback=True)
    else:
        # Текстовый ответ
        answer = update.message.text
        return await save_answer_and_continue(update, context, answer)


async def save_answer_and_continue(update: Update, context: ContextTypes.DEFAULT_TYPE, answer: str, is_callback: bool = False) -> int:
    """
    Сохраняет ответ и переходит к следующему вопросу
    """
    doc_type = context.user_data['document_type']
    doc_subtype = context.user_data['document_subtype']
    current_question = context.user_data.get('current_question', 0)
    
    # Получаем вопросы для сохранения ключа
    questions = get_questions_for_document(doc_type, doc_subtype)
    
    if current_question < len(questions):
        question_key = questions[current_question]['key']
        context.user_data['document_data'][question_key] = answer
        
        # Подтверждение сохранения ответа
        confirmation_text = f"✅ Ответ сохранен: *{answer}*"
        
        if is_callback:
            await update.callback_query.message.reply_text(
                confirmation_text,
                parse_mode='Markdown'
            )
        else:
            await update.message.reply_text(
                confirmation_text,
                parse_mode='Markdown'
            )
    
    # Переходим к следующему вопросу
    context.user_data['current_question'] = current_question + 1
    return await ask_next_question(update, context, is_callback=is_callback)


async def start_document_generation(update: Update, context: ContextTypes.DEFAULT_TYPE, is_callback: bool = False) -> int:
    """
    Начинает процесс генерации документа
    """
    doc_type = context.user_data['document_type']
    doc_subtype = context.user_data['document_subtype']
    document_data = context.user_data['document_data']
    
    doc_info = DOCUMENT_TYPES[doc_type]
    subtype_name = doc_info["subtypes"][doc_subtype]
    
    generation_text = (
        f"🎯 **Генерация документа: {doc_info['name']} - {subtype_name}**\n\n"
        "⏳ Собранная информация обрабатывается...\n"
        "📝 Создаю профессиональный юридический документ...\n\n"
        "*Это может занять несколько секунд*"
    )
    
    if is_callback:
        await update.callback_query.message.reply_text(
            generation_text,
            parse_mode='Markdown'
        )
        # Показываем индикатор печати
        await update.callback_query.message.chat.send_action(ChatAction.TYPING)
    else:
        await update.message.reply_text(
            generation_text,
            parse_mode='Markdown'
        )
        await update.message.chat.send_action(ChatAction.TYPING)
    
    try:
        # Генерируем документ через GigaChat
        generated_document = await generate_document_with_gigachat(doc_type, doc_subtype, document_data)
        context.user_data['generated_document'] = generated_document
        # Также сохраняем полный текст для экспорта
        context.user_data['full_document_text'] = generated_document
        
        # Показываем результат
        return await show_generated_document(update, context, is_callback)
        
    except Exception as e:
        logger.error(f"Ошибка генерации документа: {e}")
        error_text = (
            "❌ **Ошибка генерации документа**\n\n"
            "Произошла техническая ошибка. Попробуйте:\n"
            "• Проверить правильность введенных данных\n"
            "• Повторить попытку через минуту\n"
            "• Обратиться в поддержку\n\n"
            "Что делаем дальше?"
        )
        
        keyboard = [
            [InlineKeyboardButton("🔄 Повторить генерацию", callback_data="retry_generation")],
            [InlineKeyboardButton("📝 Изменить данные", callback_data="edit_data")],
            [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        if is_callback:
            await update.callback_query.message.reply_text(
                error_text,
                reply_markup=reply_markup,
                parse_mode='Markdown'
            )
        else:
            await update.message.reply_text(
                error_text,
                reply_markup=reply_markup,
                parse_mode='Markdown'
            )
        
        return DocumentStates.DOCUMENT_GENERATION.value


async def generate_document_with_gigachat(doc_type: str, doc_subtype: str, document_data: Dict[str, str]) -> str:
    """
    Генерирует документ с помощью GigaChat API
    """
    # Формируем системный промпт для генерации документа
    system_prompt = get_document_generation_prompt(doc_type, doc_subtype)
    
    # Формируем пользовательский промпт с данными
    user_prompt = format_user_data_for_prompt(doc_type, doc_subtype, document_data)
    
    # Генерируем документ
    response = await gigachat_client.generate_document(system_prompt, user_prompt)
    
    return response


def get_document_generation_prompt(doc_type: str, doc_subtype: str) -> str:
    """
    Возвращает системный промпт для генерации конкретного типа документа
    """
    # Базовые требования для всех документов
    base_requirements = """
КРИТИЧЕСКИ ВАЖНО - ДОКУМЕНТ БУДЕТ ОТКЛОНЕН, ЕСЛИ:
• Менее 1000 слов
• Нет ссылок на конкретные статьи законов РФ (минимум 5 статей)
• Отсутствуют обязательные разделы
• Используются общие фразы вместо конкретных юридических формулировок
• Документ не соответствует официально-деловому стилю

ОБЯЗАТЕЛЬНЫЕ ТРЕБОВАНИЯ КО ВСЕМ ДОКУМЕНТАМ:
• Юридически точные формулировки
• Ссылки на статьи Гражданского кодекса РФ, АПК РФ, ГПК РФ или других применимых законов
• Структурированность и официально-деловой стиль
• Минимум 1000-1500 слов для полноценного документа
• Все разделы должны быть детально проработаны
"""

    prompts = {
        "contract": {
            "rent": f"""{base_requirements}

Создайте ПОДРОБНЫЙ договор аренды, который будет полностью соответствовать российскому законодательству.

СПЕЦИФИЧЕСКИЕ ТРЕБОВАНИЯ ДЛЯ ДОГОВОРА АРЕНДЫ:
• Обязательно укажите данные сторон (ФИО, паспортные данные, адреса)
• Детально опишите предмет аренды со всеми характеристиками
• Определите права и обязанности сторон со ссылками на ст. 611, 615 ГК РФ
• Укажите размер арендной платы и порядок внесения (ст. 614 ГК РФ)
• Пропишите сроки действия договора (ст. 610 ГК РФ)
• Детально опишите ответственность сторон
• Включите порядок разрешения споров и условия расторжения
• Ссылайтесь на статьи 606-625 ГК РФ для аренды

ОБЯЗАТЕЛЬНАЯ СТРУКТУРА:
1. Преамбула с полными данными сторон
2. Предмет договора (ст. 606 ГК РФ) - детальное описание
3. Права и обязанности арендодателя (ст. 611 ГК РФ)
4. Права и обязанности арендатора (ст. 615 ГК РФ)
5. Арендная плата и порядок расчетов (ст. 614 ГК РФ)
6. Срок действия договора (ст. 610 ГК РФ)
7. Передача и возврат имущества
8. Техническое состояние и ремонт
9. Коммунальные и эксплуатационные расходы
10. Ответственность сторон и штрафные санкции
11. Форс-мажорные обстоятельства
12. Порядок изменения и расторжения договора
13. Разрешение споров
14. Заключительные положения
15. Реквизиты и подписи сторон

ДОПОЛНИТЕЛЬНЫЕ РАЗДЕЛЫ:
• Страхование арендованного имущества
• Порядок проведения ремонта
• Правила пользования общим имуществом
• Ответственность за коммунальные платежи""",
            
            "sale": f"""{base_requirements}

Создайте ПОДРОБНЫЙ договор купли-продажи, который будет полностью соответствовать российскому законодательству.

СПЕЦИФИЧЕСКИЕ ТРЕБОВАНИЯ ДЛЯ ДОГОВОРА КУПЛИ-ПРОДАЖИ:
• Полные данные сторон (продавец и покупатель)
• Детальное описание товара с техническими характеристиками
• Цена и полный порядок расчетов (ст. 485 ГК РФ)
• Порядок передачи товара (ст. 458 ГК РФ)
• Переход права собственности (ст. 223 ГК РФ)
• Гарантии качества (ст. 470 ГК РФ)
• Ответственность за недостатки (ст. 475 ГК РФ)
• Ссылки на главу 30 ГК РФ (купля-продажа)

ОБЯЗАТЕЛЬНАЯ СТРУКТУРА:
1. Преамбула с данными сторон
2. Предмет договора - детальное описание товара
3. Цена товара и порядок оплаты (ст. 485 ГК РФ)
4. Качество товара (ст. 469 ГК РФ)
5. Порядок передачи товара (ст. 458 ГК РФ)
6. Переход права собственности и риска (ст. 223, 224 ГК РФ)
7. Права и обязанности продавца
8. Права и обязанности покупателя
9. Гарантии и ответственность за недостатки
10. Ответственность сторон
11. Форс-мажор
12. Разрешение споров
13. Заключительные положения
14. Реквизиты и подписи"""
        },
        "lawsuit": {
            "debt": f"""{base_requirements}

Создайте ПОДРОБНОЕ исковое заявление о взыскании долга, которое будет полностью соответствовать российскому процессуальному законодательству.

СПЕЦИФИЧЕСКИЕ ТРЕБОВАНИЯ ДЛЯ ИСКОВОГО ЗАЯВЛЕНИЯ:
• Строгое соблюдение требований ст. 131 ГПК РФ или ст. 125 АПК РФ
• Полные данные истца и ответчика
• Детальное изложение обстоятельств дела
• Правовое обоснование требований со ссылками на законы
• Точный расчет взыскиваемой суммы
• Перечень доказательств
• Ссылки на ст. 309, 310, 395 ГК РФ для обязательств

ОБЯЗАТЕЛЬНАЯ СТРУКТУРА ПО СТ. 131 ГПК РФ:
1. Наименование суда
2. Данные истца: ФИО, место жительства, контакты
3. Данные ответчика: ФИО, место жительства
4. Цена иска (если подлежит оценке)
5. Обстоятельства дела:
   - Основание возникновения долга
   - Размер задолженности
   - Попытки досудебного урегулирования
   - Нарушение обязательств ответчиком
6. Правовое обоснование:
   - Ссылки на ст. 309 ГК РФ (общие положения об обязательствах)
   - Ст. 310 ГК РФ (недопустимость одностороннего отказа)
   - Ст. 395 ГК РФ (ответственность за неисполнение денежного обязательства)
7. Расчет суммы к взысканию:
   - Основной долг
   - Проценты/пени
   - Госпошлина
8. Доказательства (договоры, расписки, переписка)
9. Перечень прилагаемых документов
10. Дата и подпись истца

ИСКОВЫЕ ТРЕБОВАНИЯ должны быть сформулированы четко и конкретно.""",
            
            "payment": f"""{base_requirements}

Создайте ПОДРОБНОЕ исковое заявление о взыскании задолженности по оплате, соответствующее требованиям процессуального законодательства.

Применяйте те же требования, что и для взыскания долга, но с акцентом на:
• Договорные отношения по оплате услуг/товаров
• Ст. 516 ГК РФ (оплата товаров)
• Ст. 781 ГК РФ (оплата услуг)
• Детальный расчет неустойки и процентов"""
        },
        "claim": {
            "payment": f"""{base_requirements}

Создайте ПОДРОБНУЮ досудебную претензию о взыскании долга, соответствующую требованиям российского законодательства.

СПЕЦИФИЧЕСКИЕ ТРЕБОВАНИЯ ДЛЯ ПРЕТЕНЗИИ:
• Соблюдение досудебного порядка урегулирования
• Четкое изложение нарушений
• Конкретные требования с расчетом
• Разумный срок для ответа (10-30 дней)
• Предупреждение о судебных последствиях
• Ссылки на ст. 309 ГК РФ и договорные условия

ОБЯЗАТЕЛЬНАЯ СТРУКТУРА:
1. Данные отправителя (кредитора):
   - ФИО/наименование организации
   - Адрес, контактные данные
2. Данные получателя (должника)
3. Основание претензии:
   - Реквизиты договора/обязательства
   - Условия исполнения
   - Факт нарушения обязательств
4. Правовое обоснование:
   - Ссылка на ст. 309 ГК РФ (принцип надлежащего исполнения)
   - Ст. 310 ГК РФ (недопустимость одностороннего отказа)
   - Договорные условия
5. Требования:
   - Погашение основного долга
   - Уплата процентов/неустойки (ст. 395 ГК РФ)
   - Возмещение расходов
6. Расчет задолженности:
   - Основная сумма
   - Проценты за пользование чужими средствами
   - Неустойка (если предусмотрена)
7. Срок исполнения требований (10-30 дней)
8. Предупреждение о последствиях:
   - Обращение в суд
   - Взыскание судебных расходов
   - Принудительное исполнение
9. Перечень прилагаемых документов
10. Дата и подпись

ДОПОЛНИТЕЛЬНЫЕ ЭЛЕМЕНТЫ:
• Ссылка на принцип добросовестности (ст. 10 ГК РФ)
• Указание на попытки досудебного урегулирования
• Расчет будущих судебных расходов"""
        }
    }
    
    # Возвращаем промпт для конкретного типа или улучшенный базовый
    if doc_type in prompts and doc_subtype in prompts[doc_type]:
        return prompts[doc_type][doc_subtype]
    
    # Улучшенный базовый промпт для других типов
    return f"""{base_requirements}

Создайте профессиональный юридический документ, который будет полностью соответствовать российскому законодательству.

ТРЕБОВАНИЯ К ДОКУМЕНТУ:
• Минимум 1000 слов
• Официально-деловой стиль
• Ссылки на применимые статьи законов РФ
• Четкая структура с пронумерованными разделами
• Детальная проработка всех аспектов
• Практическая применимость

Документ должен содержать все необходимые элементы для его использования в реальной практике."""


def format_user_data_for_prompt(doc_type: str, doc_subtype: str, document_data: Dict[str, str]) -> str:
    """
    Форматирует данные пользователя в промпт для GigaChat
    """
    # Получаем информацию о типе документа
    doc_info = DOCUMENT_TYPES.get(doc_type, {})
    subtype_info = doc_info.get('subtypes', {}).get(doc_subtype, 'документ')
    
    data_text = f"""Создайте {subtype_info} на основе следующих данных:

ИСХОДНЫЕ ДАННЫЕ ДЛЯ ДОКУМЕНТА:
"""
    
    # Форматируем данные более структурированно
    for key, value in document_data.items():
        if value and value.strip():  # Проверяем что значение не пустое
            data_text += f"• **{key}:** {value}\n"
    
    data_text += f"""

ВАЖНЫЕ УКАЗАНИЯ:
• Используйте ВСЕ предоставленные данные в документе
• Где данные отсутствуют, укажите стандартные формулировки с пометкой [указать при заполнении]
• Обязательно включите все разделы из требуемой структуры
• Документ должен быть готов к использованию после подстановки конкретных данных
• Минимальный объем: 1000-1500 слов
• Обязательно включите минимум 5 ссылок на статьи российских законов

РЕЗУЛЬТАТ: Создайте полноценный, детальный и юридически корректный документ, готовый к практическому применению."""
    
    return data_text


async def show_generated_document(update: Update, context: ContextTypes.DEFAULT_TYPE, is_callback: bool = False) -> int:
    """
    Показывает сгенерированный документ пользователю с умной разбивкой
    """
    generated_document = context.user_data['generated_document']
    
    # Сохраняем полный документ для экспорта в Word
    context.user_data['full_document_text'] = generated_document
    
    # Формируем заголовок для документа
    doc_type = context.user_data.get('document_type', '')
    doc_subtype = context.user_data.get('document_subtype', '')
    doc_info = DOCUMENT_TYPES.get(doc_type, {})
    subtype_name = doc_info.get('subtypes', {}).get(doc_subtype, 'документ')
    
    header_text = (
        f"📄 **Ваш документ готов!**\n\n"
        f"📋 **Тип:** {doc_info.get('name', 'Документ')} - {subtype_name}\n"
        f"📏 **Размер:** {len(generated_document)} символов\n\n"
        "Проверьте документ и выберите действие:"
    )
    
    # Создаем клавиатуру для действий с документом
    keyboard = [
        [InlineKeyboardButton("✅ Принять документ", callback_data="accept_document")],
        [InlineKeyboardButton("✏️ Изменить", callback_data="edit_document")],
        [InlineKeyboardButton("➕ Дополнить", callback_data="supplement_document")],
        [InlineKeyboardButton("🔄 Перегенерировать", callback_data="regenerate_document")],
        [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    # Формируем полный текст для отправки
    full_message = f"{header_text}\n\n```\n{generated_document}\n```"
    
    try:
        # Используем умную разбивку из analysis_handlers
        if is_callback:
            await send_long_message(
                update.callback_query.message.chat,
                full_message,
                reply_markup=reply_markup,
                parse_mode='Markdown'
            )
        else:
            await send_long_message(
                update.message.chat,
                full_message,
                reply_markup=reply_markup,
                parse_mode='Markdown'
            )
    except Exception as e:
        # Если не удалось отправить с Markdown, пробуем без него
        logger.warning(f"Ошибка отправки с Markdown: {e}")
        
        # Формируем текст без Markdown
        plain_message = f"{header_text}\n\n{generated_document}"
        
        try:
            if is_callback:
                await send_long_message(
                    update.callback_query.message.chat,
                    plain_message,
                    reply_markup=reply_markup
                )
            else:
                await send_long_message(
                    update.message.chat,
                    plain_message,
                    reply_markup=reply_markup
                )
        except Exception as e2:
            # Последняя попытка - обрезанная версия
            logger.error(f"Ошибка отправки документа: {e2}")
            
            max_length = 3500
            truncated_document = generated_document[:max_length] + "\n\n... (документ обрезан для отображения)"
            
            result_text = (
                f"📄 **Ваш документ готов!**\n\n"
                f"{truncated_document}\n\n"
                "⚠️ Документ слишком большой для полного отображения.\n"
                "Полная версия будет доступна в Word файле.\n\n"
                "Проверьте документ и выберите действие:"
            )
            
            if is_callback:
                await update.callback_query.message.reply_text(
                    result_text,
                    reply_markup=reply_markup
                )
            else:
                await update.message.reply_text(
                    result_text,
                    reply_markup=reply_markup
                )
    
    return DocumentStates.DOCUMENT_REVIEW.value


async def handle_document_review(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработчик действий с готовым документом
    """
    query = update.callback_query
    await query.answer()
    
    if query.data == "accept_document":
        return await finalize_document(update, context)
    elif query.data == "edit_document":
        return await start_document_editing(update, context)
    elif query.data == "supplement_document":
        return await start_document_supplement(update, context)
    elif query.data == "regenerate_document":
        return await regenerate_document(update, context)
    
    return DocumentStates.DOCUMENT_REVIEW.value


async def start_document_editing(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Начало редактирования документа
    """
    query = update.callback_query
    
    edit_text = (
        "✏️ **Редактирование документа**\n\n"
        "Опишите, что нужно изменить в документе:\n"
        "• Исправить ошибки\n"
        "• Изменить формулировки\n"
        "• Добавить/убрать разделы\n"
        "• Изменить данные\n\n"
        "Напишите ваши пожелания:"
    )
    
    keyboard = [
        [InlineKeyboardButton("❌ Отмена", callback_data="cancel_edit")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        edit_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return DocumentStates.DOCUMENT_EDITING.value


async def start_document_supplement(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Начало дополнения документа
    """
    query = update.callback_query
    
    supplement_text = (
        "➕ **Дополнение документа**\n\n"
        "Укажите, что нужно добавить к документу:\n"
        "• Дополнительные условия\n"
        "• Новые разделы\n"
        "• Приложения\n"
        "• Специальные оговорки\n\n"
        "Опишите ваши дополнения:"
    )
    
    keyboard = [
        [InlineKeyboardButton("❌ Отмена", callback_data="cancel_supplement")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        supplement_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return DocumentStates.DOCUMENT_EDITING.value


async def process_document_changes(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработка изменений документа
    """
    user_changes = update.message.text
    
    # Показываем индикатор обработки
    await update.message.chat.send_action(ChatAction.TYPING)
    
    processing_text = (
        "🔄 **Обработка изменений...**\n\n"
        f"📝 Ваши пожелания: *{user_changes}*\n\n"
        "⏳ Применяю изменения к документу..."
    )
    
    await update.message.reply_text(
        processing_text,
        parse_mode='Markdown'
    )
    
    try:
        # Получаем текущий документ
        current_document = context.user_data['generated_document']
        
        # Формируем промпт для изменений
        edit_prompt = f"""
Исходный документ:
{current_document}

Требуемые изменения:
{user_changes}

Внеси указанные изменения в документ, сохранив его юридическую корректность и структуру.
"""
        
        # Генерируем обновленный документ
        system_prompt = "Ты профессиональный юрист. Внеси изменения в документ согласно пожеланиям пользователя, сохранив юридическую корректность и профессиональную структуру."
        
        updated_document = await generate_document_with_gigachat("edit", "changes", {"edit_prompt": edit_prompt})
        
        # Сохраняем обновленный документ
        context.user_data['generated_document'] = updated_document
        # Также сохраняем полный текст для экспорта
        context.user_data['full_document_text'] = updated_document
        
        success_text = (
            "✅ **Изменения применены!**\n\n"
            "Документ обновлен согласно вашим пожеланиям.\n"
            "Проверьте результат и выберите действие:"
        )
        
        await update.message.reply_text(success_text, parse_mode='Markdown')
        
        # Показываем обновленный документ
        return await show_generated_document(update, context)
        
    except Exception as e:
        logger.error(f"Ошибка при редактировании документа: {e}")
        
        error_text = (
            "❌ **Ошибка редактирования**\n\n"
            "Не удалось применить изменения.\n"
            "Попробуйте:\n"
            "• Уточнить пожелания\n"
            "• Упростить требования\n"
            "• Повторить попытку\n\n"
            "Что делаем?"
        )
        
        keyboard = [
            [InlineKeyboardButton("🔄 Попробовать снова", callback_data="edit_document")],
            [InlineKeyboardButton("⬅️ К документу", callback_data="show_document")],
            [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            error_text,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
        return DocumentStates.DOCUMENT_REVIEW.value


async def regenerate_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Перегенерация документа
    """
    query = update.callback_query
    
    await query.message.reply_text(
        "🔄 **Перегенерация документа...**\n\n"
        "⏳ Создаю новую версию документа с теми же данными...",
        parse_mode='Markdown'
    )
    
    # Перегенерируем с теми же данными
    return await start_document_generation(update, context, is_callback=True)


async def finalize_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Финализация документа - создание DOCX и сохранение
    """
    query = update.callback_query
    
    finalization_text = (
        "💾 **Финализация документа...**\n\n"
        "⏳ Создаю файл Word (.docx)...\n"
        "📤 Подготавливаю к загрузке...\n\n"
        "*Это займет несколько секунд*"
    )
    
    await query.message.reply_text(
        finalization_text,
        parse_mode='Markdown'
    )
    
    try:
        # Используем полный текст документа для экспорта (не разбитый)
        full_document = context.user_data.get('full_document_text', context.user_data['generated_document'])
        
        # Создаем DOCX файл с полным текстом
        docx_file_path = await create_docx_document(full_document, context.user_data)
        
        # Отправляем файл пользователю
        with open(docx_file_path, 'rb') as doc_file:
            await query.message.reply_document(
                document=doc_file,
                filename=get_document_filename(context.user_data),
                caption=(
                    "📄 **Ваш документ готов!**\n\n"
                    "✅ Файл создан в формате Word (.docx)\n"
                    "📁 Можете скачать и использовать\n"
                    "📏 Полный документ сохранен целиком\n\n"
                    "🔍 Обязательно проверьте документ перед использованием!"
                ),
                parse_mode='Markdown'
            )
        
        # Удаляем временный файл
        import os
        try:
            os.remove(docx_file_path)
        except Exception:
            pass
        
        # Запрашиваем оценку
        return await request_document_rating(update, context)
        
    except Exception as e:
        logger.error(f"Ошибка создания DOCX: {e}")
        
        error_text = (
            "❌ **Ошибка создания файла**\n\n"
            "Не удалось создать DOCX файл.\n"
            "Документ остается доступен в чате.\n\n"
            "Попробуйте:\n"
            "• Скопировать текст документа\n"
            "• Повторить создание файла\n"
            "• Обратиться в поддержку"
        )
        
        keyboard = [
            [InlineKeyboardButton("🔄 Повторить создание", callback_data="retry_docx")],
            [InlineKeyboardButton("📋 Показать текст", callback_data="show_document")],
            [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.message.reply_text(
            error_text,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        
        return DocumentStates.DOCUMENT_FINALIZATION.value


async def create_docx_document(document_text: str, user_data: Dict[str, Any]) -> str:
    """
    Создает DOCX файл из текста документа
    """
    try:
        from docx import Document
        from docx.shared import Inches
        import os
        import tempfile
        
        # Создаем новый документ Word
        doc = Document()
        
        # Устанавливаем поля
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # Добавляем текст документа
        # Разбиваем на абзацы
        paragraphs = document_text.split('\n\n')
        
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                # Проверяем, является ли это заголовком
                if (paragraph_text.strip().isupper() or 
                    paragraph_text.strip().startswith('ДОГОВОР') or
                    paragraph_text.strip().startswith('ИСКОВОЕ') or
                    paragraph_text.strip().startswith('ПРЕТЕНЗИЯ')):
                    # Добавляем как заголовок
                    heading = doc.add_heading(paragraph_text.strip(), level=1)
                    heading.alignment = 1  # Центрирование
                else:
                    # Добавляем как обычный абзац
                    p = doc.add_paragraph(paragraph_text.strip())
                    p.alignment = 3  # Выравнивание по ширине
        
        # Создаем временный файл
        temp_dir = tempfile.gettempdir()
        filename = get_document_filename(user_data)
        file_path = os.path.join(temp_dir, filename)
        
        # Сохраняем документ
        doc.save(file_path)
        
        return file_path
        
    except ImportError:
        # Если библиотека python-docx не установлена
        raise Exception("Библиотека python-docx не найдена")
    except Exception as e:
        raise Exception(f"Ошибка создания DOCX: {e}")


def get_document_filename(user_data: Dict[str, Any]) -> str:
    """
    Генерирует имя файла для документа
    """
    doc_type = user_data.get('document_type', 'document')
    doc_subtype = user_data.get('document_subtype', 'general')
    
    # Переводим типы в читаемые имена
    type_names = {
        'contract': 'Договор',
        'lawsuit': 'Исковое_заявление',
        'claim': 'Претензия',
        'agreement': 'Соглашение',
        'power_of_attorney': 'Доверенность',
        'application': 'Заявление',
        'protocol': 'Протокол',
        'act': 'Акт'
    }
    
    subtype_names = {
        'rent': 'аренды',
        'sale': 'купли_продажи',
        'debt': 'взыскание_долга',
        'payment': 'взыскание_долга',
        'general': 'общий'
    }
    
    type_name = type_names.get(doc_type, 'Документ')
    subtype_name = subtype_names.get(doc_subtype, doc_subtype)
    
    import datetime
    date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    
    return f"{type_name}_{subtype_name}_{date_str}.docx"


async def request_document_rating(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Запрос оценки документа от пользователя
    """
    rating_text = (
        "⭐ **Оцените качество документа**\n\n"
        "Помогите нам улучшить сервис!\n"
        "Оцените созданный документ от 1 до 5:"
    )
    
    keyboard = []
    for i in range(1, 6):
        stars = "⭐" * i
        keyboard.append([InlineKeyboardButton(f"{stars} {i}", callback_data=f"rate_{i}")])
    
    keyboard.append([InlineKeyboardButton("➡️ Пропустить оценку", callback_data="skip_rating")])
    keyboard.append([InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.callback_query.message.reply_text(
        rating_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    return DocumentStates.DOCUMENT_FINALIZATION.value


async def handle_document_rating(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработка оценки документа
    """
    query = update.callback_query
    await query.answer()
    
    if query.data.startswith("rate_"):
        rating = int(query.data.replace("rate_", ""))
        
        # TODO: Сохранить оценку в Supabase
        # await save_document_rating(context.user_data, rating)
        
        thank_you_text = (
            f"🙏 **Спасибо за оценку: {rating} ⭐**\n\n"
            "Ваша обратная связь поможет нам улучшить качество документов!\n\n"
            "📄 **Что еще могу для вас сделать?**"
        )
        
        keyboard = [
            [InlineKeyboardButton("📝 Создать новый документ", callback_data="new_document")],
            [InlineKeyboardButton("📋 Анализ документа", callback_data="analyze")],
            [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
        ]
        
    else:  # skip_rating
        thank_you_text = (
            "👍 **Понятно, спасибо!**\n\n"
            "📄 **Что еще могу для вас сделать?**"
        )
        
        keyboard = [
            [InlineKeyboardButton("📝 Создать новый документ", callback_data="new_document")],
            [InlineKeyboardButton("📋 Анализ документа", callback_data="analyze")],
            [InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu")]
        ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        thank_you_text,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )
    
    # Очищаем данные документа
    context.user_data.pop('document_data', None)
    context.user_data.pop('document_type', None)
    context.user_data.pop('document_subtype', None)
    context.user_data.pop('generated_document', None)
    
    return ConversationHandler.END